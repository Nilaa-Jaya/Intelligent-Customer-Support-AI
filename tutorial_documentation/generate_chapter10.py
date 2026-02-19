"""
Generate Chapter 10: Knowledge Base Retrieval Agent
Concise but complete — covers every gap from the original.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ──────────────────────────────────────────────────────────────────

def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    run = hd.runs[0] if hd.runs else hd.add_run(text)
    colours = {1: RGBColor(0x1F,0x49,0x7D), 2: RGBColor(0x2E,0x74,0xB5), 3: RGBColor(0x1F,0x49,0x7D)}
    sizes   = {1: 18, 2: 14, 3: 12}
    run.font.color.rgb = colours.get(level, RGBColor(0,0,0))
    run.font.size = Pt(sizes.get(level, 11))
    return hd

def p(doc, text, bold=False, italic=False, size=11, colour=None):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    return para

def code(doc, text, caption=""):
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_after = Pt(0)
        r = cap.add_run(f"# {caption}")
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55,0x55,0x55)

    para = doc.add_paragraph()
    para.style = doc.styles['No Spacing']
    para.paragraph_format.left_indent  = Inches(0.25)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(8)

    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4F8')
    pPr.append(shd)

    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    return para

def callout(doc, text, kind="KEY INSIGHT"):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(6)
    label = para.add_run(f"{kind}:  ")
    label.bold = True
    label.font.color.rgb = RGBColor(0x1F,0x49,0x7D)
    body = para.add_run(text)
    body.italic = True
    body.font.size = Pt(10)

def bullet(doc, text, level=0):
    b = doc.add_paragraph(style='List Bullet')
    b.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    b.add_run(text).font.size = Pt(11)


# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

# Title
title = doc.add_heading('Chapter 10: Knowledge Base Retrieval Agent', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('SmartSupport AI — Complete Reference')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
sub.runs[0].font.size = Pt(12)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 10.1  WHAT IS THIS AGENT AND WHY DOES IT EXIST?
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.1  What Is the KB Retrieval Agent?', 2)

p(doc,
  'The KB Retrieval Agent is the third node in the LangGraph workflow. '
  'Its job is simple: before any specialist agent (Technical, Billing, etc.) '
  'writes a response, this agent searches a database of 30 FAQ articles and '
  'attaches the most relevant ones to the state. The specialist agent then '
  'uses those articles to ground its answer in real, company-specific facts '
  'instead of relying on the LLM\'s general training data.')

p(doc,
  'Without this step, the LLM might invent plausible-sounding but incorrect '
  'policy information. With it, the LLM is told "here are the actual company '
  'FAQs — use these." This pattern is called RAG (Retrieval-Augmented Generation).')

callout(doc,
  'The KB agent does NOT generate any text. It only reads from the vector '
  'store and writes to state["kb_results"]. Generating the response is always '
  'done by the downstream specialist agent.')

# ─────────────────────────────────────────────────────────────────────────────
# 10.2  COMPLETE DATA FLOW
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.2  Complete Data Flow (End to End)', 2)

p(doc, 'Five distinct steps happen in sequence every time a query arrives:')

code(doc, '''\
STEP 1 — App startup (runs ONCE, not per query)
  faqs.json on disk
      |
      v  load_faqs_from_json()
  List of dicts  [{"id":1, "category":"Technical",
                   "question":"Why does my app crash?",
                   "answer":"...",
                   "text":"Q: Why... A: ..."}]   ← Q+A combined here
      |
      v  VectorStore.add_documents()
  SentenceTransformer encodes every "text" field → 384-dim float32 vector
      |
      v  faiss.IndexFlatL2.add()
  Vectors stored in FAISS in-memory index
      |
      v  VectorStore.save()
  faiss_index.index  +  metadata.json   saved to data/knowledge_base/

──────────────────────────────────────────────────────────────────────
STEP 2 — Every subsequent startup (loads in milliseconds)
  VectorStore.load()  →  reads .index file + metadata.json  →  ready

──────────────────────────────────────────────────────────────────────
STEP 3 — retrieve_from_kb() called by LangGraph (ONCE per query)
  state["query"]    = "My app keeps crashing"
  state["category"] = "Technical"

──────────────────────────────────────────────────────────────────────
STEP 4 — Two-step filtering inside KnowledgeBaseRetriever.retrieve()
  VectorStore.search()
      Encode query → query vector
      FAISS search: fetch top k*3 = 9 candidates   ← over-fetch for filter
      Keep only category == "Technical"             ← filter 1: category
      Return up to k=3 results
  KnowledgeBaseRetriever.retrieve()
      Discard results with score < 0.3              ← filter 2: min_score
      Return final list

──────────────────────────────────────────────────────────────────────
STEP 5 — Results stored in state, used by specialist agent
  state["kb_results"] = [
      {"title": "Why does my app crash?",
       "content": "App crashes can be caused by ...",
       "category": "Technical",
       "score": 0.73},
      ...
  ]
  → Technical agent injects these into its LLM prompt
''', 'Full pipeline from disk file to LLM prompt')

# ─────────────────────────────────────────────────────────────────────────────
# 10.3  THE FAQ DATA (the raw material)
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.3  The FAQ Data Structure', 2)

p(doc,
  'Everything starts with data/knowledge_base/faqs.json — 30 FAQ entries '
  'spread across four categories: Technical (10), Billing (8), '
  'Account (7), General (5). Each entry has this shape:')

code(doc, '''\
{
  "faqs": [
    {
      "id": 1,
      "category": "Technical",
      "question": "Why does my app keep crashing?",
      "answer": "App crashes can be caused by several factors:
                 1) Outdated app version - Update to the latest version.
                 2) Low device storage - Clear cache and free up 500MB.
                 3) Conflicting apps - Close background apps.
                 ..."
    },
    {
      "id": 12,
      "category": "Billing",
      "question": "How do I get a refund?",
      "answer": "Refund requests can be submitted within 30 days ..."
    }
    // ... 28 more entries
  ]
}
''', 'data/knowledge_base/faqs.json — one entry per category')

# ─────────────────────────────────────────────────────────────────────────────
# 10.4  THE SENTENCE TRANSFORMER (the "brain" of retrieval)
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.4  The Sentence Transformer Model', 2)

p(doc,
  'A Sentence Transformer is a neural network specifically trained to convert '
  'a sentence into a fixed-size vector of numbers such that sentences with '
  'similar meaning end up close together in vector space. Unlike a regular LLM '
  'that generates text token by token, a Sentence Transformer only encodes — '
  'it reads a sentence and outputs one vector. No text generation happens here.')

p(doc,
  'SmartSupport AI uses all-MiniLM-L6-v2. Here is why this model was chosen:')

bullet(doc, 'Size: 22 million parameters — tiny compared to LLMs (70 billion). '
       'Loads in ~1 second, uses ~90 MB RAM.')
bullet(doc, 'Speed: Encodes a sentence in milliseconds on CPU — no GPU needed.')
bullet(doc, 'Output: 384-dimensional vector — small enough for fast FAISS search.')
bullet(doc, 'Quality: Strong semantic accuracy on sentence-similarity benchmarks, '
       'good enough for FAQ retrieval where the vocabulary is limited and domain-specific.')
bullet(doc, 'Language: English-only, which matches the FAQ content.')

code(doc, '''\
# What "384 dimensions" means in practice:
# Each FAQ is represented as a list of 384 floating-point numbers.
# Example (truncated — real vector has 384 values):
"Why does my app keep crashing?"  →  [0.12, -0.34, 0.56, 0.01, ..., 0.89]
"My application crashes on startup" → [0.11, -0.31, 0.58, 0.03, ..., 0.87]

# These two vectors are CLOSE because the sentences mean the same thing.
# The model learned this from millions of example sentence pairs during training.

# A completely different sentence would be far away:
"What is your refund policy?"     → [-0.45, 0.72, -0.21, 0.65, ..., -0.33]
''', 'How the 384-dimensional vector captures meaning')

callout(doc,
  'The model is loaded ONCE via the get_kb_retriever() singleton. '
  'Loading it repeatedly for each query would add ~1 second of delay per request.')

# ─────────────────────────────────────────────────────────────────────────────
# 10.5  load_faqs_from_json() — THE ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.5  load_faqs_from_json() — Preparing Documents for Indexing', 2)

p(doc,
  'This standalone function reads faqs.json and transforms each raw FAQ '
  'into the format VectorStore expects. The most important thing it does '
  'is create a combined "text" field:')

code(doc, '''\
# src/knowledge_base/vector_store.py

def load_faqs_from_json(file_path: str) -> List[Dict]:
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    documents = []
    for faq in data.get("faqs", []):

        # ← THE KEY DESIGN DECISION
        text = f"Q: {faq['question']}\\nA: {faq['answer']}"

        doc = {
            "id":       faq["id"],
            "category": faq["category"],
            "question": faq["question"],   # shown in API response (title)
            "answer":   faq["answer"],     # shown in API response (content)
            "text":     text,              # embedded by SentenceTransformer
        }
        documents.append(doc)

    return documents
''', 'src/knowledge_base/vector_store.py — load_faqs_from_json()')

p(doc,
  'Why combine Q and A into one "text" field for embedding?')

bullet(doc,
  'If only the question is embedded: a user query "how do I export my files?" '
  'might miss FAQ #4 whose question is "How do I export my data?" because '
  '"files" vs "data" differ slightly.')
bullet(doc,
  'If Q+A are combined: the answer text "Go to Settings > Data Export and '
  'select CSV/JSON/PDF" also participates in the embedding. Now a user query '
  '"where is the export button?" matches on the answer content, not just the question.')
bullet(doc, 'Result: significantly better retrieval recall, especially for '
  'queries that describe a symptom rather than ask the exact question.')

# ─────────────────────────────────────────────────────────────────────────────
# 10.6  VectorStore CLASS — all methods
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.6  VectorStore — All Methods Explained', 2)

p(doc,
  'VectorStore wraps FAISS and the SentenceTransformer into one object. '
  'It has six methods. Here is each one:')

h(doc, '__init__ — Initialisation', 3)
code(doc, '''\
def __init__(self,
             model_name="all-MiniLM-L6-v2",
             index_path="./data/knowledge_base/faiss_index",
             metadata_path="./data/knowledge_base/metadata.json"):

    self.encoder       = SentenceTransformer(model_name)  # load the model
    self.embedding_dim = self.encoder.get_sentence_embedding_dimension()  # = 384
    self.index         = None          # FAISS index (None until add_documents)
    self.documents     = []            # parallel list of raw dicts

    self.load()   # try to load existing index from disk immediately
''', '__init__: two files on disk, two data structures in memory')

p(doc, 'Note that index_path and metadata_path are two separate files. '
  'FAISS stores the vectors in a binary .index file; the document '
  'metadata (question, answer, category) lives in metadata.json because '
  'FAISS only stores numbers, not text.')

h(doc, 'add_documents — Build the Index', 3)
code(doc, '''\
def add_documents(self, documents: List[Dict]) -> None:
    texts = [doc.get("text", "") for doc in documents]  # extract the Q+A strings

    # Encode all texts at once (batch processing is faster than one by one)
    embeddings = self.encoder.encode(texts, show_progress_bar=True)
    embeddings = np.array(embeddings).astype("float32")  # FAISS requires float32

    if self.index is None:
        # IndexFlatL2: exact search using Euclidean (L2) distance.
        # "Exact" = checks every stored vector — accurate but O(n) per query.
        # Acceptable here because n = 30 FAQs (would use IndexIVFFlat for n > 10,000)
        self.index = faiss.IndexFlatL2(self.embedding_dim)  # 384

    self.index.add(embeddings)       # add vectors to FAISS
    self.documents.extend(documents) # keep the metadata in parallel
''', 'add_documents: encode text → store in FAISS')

h(doc, 'search — Similarity Search (with two built-in guards)', 3)
code(doc, '''\
def search(self, query: str, k: int = 3,
           category_filter: Optional[str] = None) -> List[Dict]:

    if self.index is None or len(self.documents) == 0:
        return []   # guard: nothing indexed yet

    # Encode the query (same model, same vector space as the FAQs)
    query_vec = self.encoder.encode([query])
    query_vec = np.array(query_vec).astype("float32")

    # Over-fetch so we have enough candidates after category filtering.
    # If category_filter="Technical", many candidates will be Billing/General
    # and will be discarded. Fetching k*3=9 ensures we likely get k=3 matches.
    search_k = k * 3 if category_filter else k
    search_k = min(search_k, len(self.documents))  # guard: don\'t ask for more than we have

    distances, indices = self.index.search(query_vec, search_k)
    # distances[0]: L2 distances (lower = more similar)
    # indices[0]:   positions in self.documents list

    results = []
    for distance, idx in zip(distances[0], indices[0]):
        if idx < len(self.documents):
            doc = self.documents[idx].copy()

            # Convert L2 distance → similarity score in range (0, 1]
            # distance=0  → score=1.0 (identical)
            # distance=1  → score=0.5
            # distance=9  → score=0.1 (barely related)
            doc["similarity_score"] = float(1 / (1 + distance))

            if category_filter is None or doc.get("category") == category_filter:
                results.append(doc)

            if len(results) >= k:
                break   # stop once we have k results

    return results
''', 'search: FAISS lookup + category filter + score conversion')

h(doc, 'save and load — Persistence', 3)
code(doc, '''\
def save(self) -> None:
    """Write index + metadata to disk after building."""
    os.makedirs(os.path.dirname(self.index_path), exist_ok=True)

    faiss.write_index(self.index, f"{self.index_path}.index")  # binary file
    with open(self.metadata_path, "w") as f:
        json.dump(self.documents, f, indent=2)                 # human-readable JSON


def load(self) -> bool:
    """Read index + metadata from disk. Called in __init__ automatically."""
    index_file = f"{self.index_path}.index"

    if not os.path.exists(index_file) or not os.path.exists(self.metadata_path):
        return False   # first run — no files yet, start fresh

    self.index     = faiss.read_index(index_file)
    with open(self.metadata_path, "r") as f:
        self.documents = json.load(f)
    return True
''', 'save / load: two files, always in sync')

p(doc,
  'After the first run, the two files exist on disk. Every subsequent '
  'startup skips re-encoding and re-indexing — load() reads the saved '
  'binary directly in milliseconds. This matters because encoding 30 FAQs '
  'takes about 2-3 seconds; loading the saved index takes under 50 ms.')

h(doc, 'clear and get_stats — Utility Methods', 3)
code(doc, '''\
def clear(self) -> None:
    """Reset to empty state (used by reload_faqs() in the retriever)."""
    self.index     = None
    self.documents = []

def get_stats(self) -> Dict:
    """Return diagnostic info — used for monitoring and debugging."""
    return {
        "total_documents":   len(self.documents),   # should be 30
        "embedding_dimension": self.embedding_dim,  # 384
        "model_name":        self.model_name,       # all-MiniLM-L6-v2
        "has_index":         self.index is not None,
    }
''', 'clear + get_stats')

# ─────────────────────────────────────────────────────────────────────────────
# 10.7  KnowledgeBaseRetriever CLASS
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.7  KnowledgeBaseRetriever — The Public Interface', 2)

p(doc,
  'KnowledgeBaseRetriever sits on top of VectorStore and is the class '
  'the rest of the application talks to. It adds two things VectorStore '
  'does not have: the min_score filter, and FAQ loading from disk.')

code(doc, '''\
# src/knowledge_base/retriever.py

class KnowledgeBaseRetriever:

    def __init__(self, faq_path="./data/knowledge_base/faqs.json",
                 model_name="all-MiniLM-L6-v2", auto_load=True):
        self.faq_path     = faq_path
        self.vector_store = VectorStore(model_name=model_name)
        if auto_load:
            self.load_faqs()   # on first use: encode + save; on restart: already saved

    # ── Loading ────────────────────────────────────────────────────────────
    def load_faqs(self) -> None:
        try:
            documents = load_faqs_from_json(self.faq_path)   # reads JSON
            if documents:
                self.vector_store.add_documents(documents)    # encode + index
                self.vector_store.save()                      # persist to disk
        except FileNotFoundError:
            pass   # graceful: KB simply won\'t return results

    # ── Core retrieval (adds the min_score filter) ─────────────────────────
    def retrieve(self, query: str, k: int = 3,
                 category: str = None, min_score: float = 0.3) -> List[Dict]:

        # Step 1: vector search + category filter (inside VectorStore)
        results = self.vector_store.search(query=query, k=k,
                                           category_filter=category)

        # Step 2: score filter (here, after vector search)
        filtered = [r for r in results if r.get("similarity_score", 0) >= min_score]
        return filtered

    # ── Maintenance ────────────────────────────────────────────────────────
    def reload_faqs(self) -> None:
        """Update FAQs without restarting the server."""
        self.vector_store.clear()
        self.load_faqs()

    def get_stats(self) -> Dict:
        stats = self.vector_store.get_stats()
        stats["faq_path"] = self.faq_path
        return stats
''', 'src/knowledge_base/retriever.py — KnowledgeBaseRetriever')

# ─────────────────────────────────────────────────────────────────────────────
# 10.8  THE SINGLETON
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.8  get_kb_retriever() — The Singleton Pattern', 2)

p(doc,
  'The agents layer never instantiates KnowledgeBaseRetriever directly. '
  'It always calls get_kb_retriever(), which returns one shared instance '
  'for the lifetime of the process:')

code(doc, '''\
# src/knowledge_base/retriever.py

_retriever: Optional[KnowledgeBaseRetriever] = None

def get_kb_retriever(faq_path="./data/knowledge_base/faqs.json",
                     force_reload=False) -> KnowledgeBaseRetriever:
    global _retriever

    if _retriever is None:
        _retriever = KnowledgeBaseRetriever(faq_path=faq_path)
    elif force_reload:
        _retriever.reload_faqs()

    return _retriever
''', 'Singleton factory in retriever.py')

p(doc, 'Why a singleton here?')
bullet(doc,
  'The SentenceTransformer model is ~90 MB and takes ~1 second to load. '
  'Creating a new retriever per query would add 1 second of latency to every '
  'customer interaction — unacceptable.')
bullet(doc,
  'The FAISS index is already in RAM once loaded. The singleton ensures '
  'it is built or loaded exactly once and reused across all concurrent requests.')
bullet(doc,
  'force_reload=True can be passed to hot-swap the FAQ content (e.g. after '
  'an admin updates faqs.json) without restarting the server.')

# ─────────────────────────────────────────────────────────────────────────────
# 10.9  THE TWO-STEP FILTER — WHY TWO SEPARATE STEPS?
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.9  The Two-Step Filtering Architecture', 2)

p(doc,
  'There are two distinct filter operations and they happen in two different '
  'classes for a deliberate reason:')

code(doc, '''\
FILTER 1 — Category (inside VectorStore.search)
  Purpose : prefer domain-relevant articles
  Mechanism: over-fetch k*3 candidates from FAISS,
             then keep only those matching state["category"]
  Why here : FAISS returns all categories mixed together;
             we need to post-filter before returning

FILTER 2 — Minimum score (inside KnowledgeBaseRetriever.retrieve)
  Purpose : discard semantically irrelevant results
  Mechanism: drop any result whose similarity_score < 0.3
  Why here : VectorStore should not know what score is "good enough"
             — that is business logic belonging to the retriever layer

Worked example (query = "slow internet connection", category = "Technical"):

  FAISS returns 9 candidates (k*3):
    score 0.71  "The app is running very slow"        category=Technical  ✓ keep
    score 0.65  "My data isn't syncing"               category=Technical  ✓ keep
    score 0.51  "How do I cancel my subscription?"    category=Billing    ✗ wrong category
    score 0.48  "How do I change my password?"        category=Account    ✗ wrong category
    score 0.44  "App keeps crashing"                  category=Technical  ✓ keep (3 found, stop)
    score 0.21  "What payment methods do you accept?" category=Billing    not reached
    ...

  After score filter (min=0.3): all three kept (0.71, 0.65, 0.44 all >= 0.3)
  Final state["kb_results"]: 3 Technical FAQs
''', 'How the two filters work together on a real example')

callout(doc,
  'If the category filter is too aggressive (e.g. only 1 relevant Technical FAQ '
  'exists), the retriever may return fewer than k results — even 0. '
  'This is intentional: an irrelevant article is worse than no article.')

# ─────────────────────────────────────────────────────────────────────────────
# 10.10  retrieve_from_kb() — THE AGENT NODE
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.10  retrieve_from_kb() — The LangGraph Node', 2)

p(doc,
  'This is the function LangGraph calls as the third node in the workflow. '
  'It is a thin wrapper that calls get_kb_retriever().retrieve() and stores '
  'the results in state["kb_results"]:')

code(doc, '''\
# src/agents/kb_retrieval.py

def retrieve_from_kb(state: AgentState) -> AgentState:
    query    = state.get("query", "")
    category = state.get("category", "General")  # set by categorize node

    try:
        kb_retriever = get_kb_retriever()

        results = kb_retriever.retrieve(
            query=query,
            k=3,
            category=category,   # e.g. "Technical"
            min_score=0.3,
        )

        # Re-shape field names to match what the specialist agents expect
        kb_results = [
            {
                "title":    r.get("question", ""),   # the FAQ question text
                "content":  r.get("answer", ""),     # the FAQ answer text
                "category": r.get("category", ""),
                "score":    r.get("similarity_score", 0.0),
            }
            for r in results
        ]

        state["kb_results"] = kb_results

        # Log top result for debugging
        if kb_results:
            top = kb_results[0]
            app_logger.info(f"Top result (score {top[\'score\']:.3f}): {top[\'title\'][:80]}")

        return state

    except Exception as e:
        app_logger.error(f"Error retrieving from KB: {e}")
        state["kb_results"] = []   # empty list → specialist agents skip KB context
        return state
''', 'src/agents/kb_retrieval.py — the node function')

# ─────────────────────────────────────────────────────────────────────────────
# 10.11  HOW DOWNSTREAM AGENTS ACTUALLY USE THE RESULTS
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.11  How Downstream Agents Use kb_results', 2)

p(doc,
  'Each specialist agent (technical, billing, account, general) follows the '
  'same pattern. Here is the Technical agent as a representative example:')

code(doc, '''\
# src/agents/technical_agent.py — inside handle_technical()

# Build kb_context string from state["kb_results"]
kb_context = ""
if state.get("kb_results"):
    kb_context = "Relevant knowledge base articles:\\n"
    for i, kb in enumerate(state["kb_results"][:2], 1):      # use top 2
        kb_context += f"{i}. {kb.get(\'title\', \'\')}:"
        kb_context += f" {kb.get(\'content\', \'\')[:200]}...\\n"  # first 200 chars

# Pass kb_context into the LLM prompt
response = llm_manager.invoke_with_retry(
    TECHNICAL_PROMPT,
    {
        "query":      state["query"],
        "sentiment":  state["sentiment"],
        "priority":   state["priority_score"],
        "context":    conversation_history_string,
        "kb_context": kb_context,    # ← injected here
    },
)
''', 'How technical_agent.py injects KB results into the LLM prompt')

p(doc, 'The prompt template itself contains the {kb_context} placeholder:')

code(doc, '''\
TECHNICAL_PROMPT = """
You are an expert technical support agent.

Customer Query: {query}
Customer Sentiment: {sentiment}
Priority Level: {priority}

{context}

{kb_context}       ← KB articles appear here if found; empty string if not

Instructions:
1. Provide a clear, step-by-step technical solution
2. Use the knowledge base articles above as your primary source
...
"""
''', 'The {kb_context} placeholder in the LLM prompt template')

p(doc,
  'When kb_results is empty (no relevant FAQ found, or KB unavailable), '
  'kb_context is an empty string and the LLM falls back to its general '
  'training knowledge. The response is still generated — it just lacks '
  'the company-specific grounding.')

callout(doc,
  'Only the top 2 KB results are injected into the prompt (not all 3) to '
  'keep the prompt within the LLM\'s context window budget. The third result '
  'is still stored in state["kb_results"] and returned in the API response '
  'as metadata for the caller to inspect.')

# ─────────────────────────────────────────────────────────────────────────────
# 10.12  QUICK REFERENCE
# ─────────────────────────────────────────────────────────────────────────────
h(doc, '10.12  Quick Reference Summary', 2)

code(doc, '''\
File                              Class / Function          Role
─────────────────────────────────────────────────────────────────────────────
data/knowledge_base/faqs.json     —                         Raw FAQ data (30 entries)
data/knowledge_base/faiss_index   —                         Saved FAISS binary index
data/knowledge_base/metadata.json —                         Saved document metadata

src/knowledge_base/vector_store.py
  load_faqs_from_json()           standalone function       Read JSON, combine Q+A
  VectorStore.__init__()          —                         Load model + call load()
  VectorStore.add_documents()     —                         Encode text → FAISS
  VectorStore.search()            —                         FAISS search + cat filter
  VectorStore.save()              —                         Write .index + .json
  VectorStore.load()              —                         Read .index + .json
  VectorStore.clear()             —                         Reset to empty
  VectorStore.get_stats()         —                         Diagnostic info

src/knowledge_base/retriever.py
  KnowledgeBaseRetriever          class                     Public interface
  KnowledgeBaseRetriever.retrieve()—                        search + min_score filter
  KnowledgeBaseRetriever.reload_faqs()—                     Hot-swap FAQs
  get_kb_retriever()              singleton factory         One instance per process

src/agents/kb_retrieval.py
  retrieve_from_kb()              LangGraph node            Calls retriever, updates state

─────────────────────────────────────────────────────────────────────────────
Key numbers
  FAQs:              30 total  (Technical 10, Billing 8, Account 7, General 5)
  Embedding model:   all-MiniLM-L6-v2  (22M params, 384-dim output)
  FAISS index type:  IndexFlatL2  (exact search, fine for n=30)
  Results returned:  k=3 (top 3 per query)
  Score threshold:   min_score=0.3  (empirically chosen)
  Over-fetch factor: k*3=9 candidates before category filter
  Prompt injection:  top 2 results only (token budget)
  State key:         state["kb_results"] — list of dicts with title/content/score
─────────────────────────────────────────────────────────────────────────────
''', 'Chapter 10 at a glance')

# ── Save ─────────────────────────────────────────────────────────────────────
out = "tutorial_documentation/SmartSupport_AI_Chapter10_KB_Retrieval.docx"
doc.save(out)
print(f"Saved: {out}")
