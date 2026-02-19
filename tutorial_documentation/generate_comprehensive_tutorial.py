"""
Generate COMPREHENSIVE Word document tutorial for SmartSupport AI
Target: 75+ pages with detailed explanations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_code_block(doc, code):
    """Add a formatted code block"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return para

def add_key_insight(doc, text):
    """Add a key insight box"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)

    run = para.add_run("KEY INSIGHT: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    run2 = para.add_run(text)
    run2.italic = True
    return para

def add_warning(doc, text):
    """Add a warning box"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)

    run = para.add_run("WARNING: ")
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)

    run2 = para.add_run(text)
    return para

def add_note(doc, text):
    """Add a note box"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    run = para.add_run("NOTE: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 150)

    run2 = para.add_run(text)
    return para

def create_table(doc, headers, rows):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)

    doc.add_paragraph()
    return table

def add_section_intro(doc, text):
    """Add section introduction paragraph"""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(12)
    return para

# ============================================================================
# PART 1: TITLE AND INTRODUCTION
# ============================================================================

def create_title_page(doc):
    """Create the title page"""
    # Main title
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_heading('SmartSupport AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Building Intelligent Customer Support Systems with AI Agents')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(70, 70, 70)

    doc.add_paragraph()

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('A Complete Guide from First Principles to Production Implementation')
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.add_run(
        "This comprehensive tutorial teaches you everything you need to know about building "
        "AI-powered customer support systems. Starting from the absolute basics of what an AI agent is, "
        "we guide you through building a complete, production-ready system with multiple specialized agents, "
        "a knowledge base with semantic search, RESTful APIs, and cloud deployment."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Key highlights box
    highlights_title = doc.add_paragraph()
    highlights_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = highlights_title.add_run('What You Will Build:')
    run.bold = True
    run.font.size = Pt(14)

    bullets = [
        '7 Specialized AI Agents orchestrated via LangGraph',
        '30 Comprehensive FAQs in FAISS Vector Store',
        'Retrieval-Augmented Generation (RAG) Pipeline',
        '8 Database Tables with Full Conversation Tracking',
        '15+ RESTful API Endpoints with Webhook Support',
        'Production-Ready Deployment with Docker and Railway',
        '38 Automated Tests with 42% Code Coverage'
    ]

    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Version 2.2.0 | January 2026 | ~75+ Pages')

    doc.add_page_break()

def create_readers_guide(doc):
    """Create the reader's guide section"""
    add_heading(doc, "Reader's Guide", 1)

    # What this document is
    add_heading(doc, "What This Document Is", 2)

    doc.add_paragraph(
        "This document is a complete, comprehensive tutorial that teaches you how to build an "
        "enterprise-grade AI-powered customer support system from scratch. Unlike quick-start guides "
        "or API documentation, this tutorial explains every concept in depth, walks through every "
        "line of code, and justifies every architectural decision."
    )

    doc.add_paragraph(
        "We start from the absolute fundamentals - explaining what an AI agent actually is, how "
        "Large Language Models work, and why we need multiple agents working together. Then we "
        "progressively build up to a complete, production-ready system."
    )

    doc.add_paragraph("By the end of this tutorial, you will:")

    outcomes = [
        "Understand what AI agents are and how they work at a fundamental level",
        "Know how to design and implement multi-agent systems",
        "Be able to build Retrieval-Augmented Generation (RAG) pipelines",
        "Understand prompt engineering and how to get consistent results from LLMs",
        "Know how to implement vector similarity search for knowledge bases",
        "Be able to build production-ready APIs with FastAPI",
        "Understand how to deploy AI applications to cloud platforms",
        "Have a complete, working customer support system you can extend"
    ]
    for outcome in outcomes:
        doc.add_paragraph(outcome, style='List Bullet')

    # Who this is for
    add_heading(doc, "Who This Tutorial Is For", 2)

    doc.add_paragraph("This tutorial is designed for several audiences:")

    audiences = [
        ("Software Developers", "who want to add AI/ML capabilities to their skill set. If you can write Python and understand basic web development, you can follow this tutorial."),
        ("Backend Engineers", "interested in building intelligent systems. The patterns here apply beyond customer support to any domain requiring AI automation."),
        ("AI/ML Engineers", "who want to learn production deployment patterns. We cover not just the ML aspects but the full software engineering lifecycle."),
        ("Computer Science Students", "building portfolio projects. This gives you a complete, impressive project that demonstrates modern AI engineering."),
        ("Technical Leads and Architects", "evaluating AI architectures. We discuss trade-offs and alternatives throughout.")
    ]

    for title, desc in audiences:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # Prerequisites
    add_heading(doc, "Prerequisites (What You Must Know)", 2)

    doc.add_paragraph(
        "We've designed this tutorial to be as accessible as possible, but some foundational "
        "knowledge is required:"
    )

    prereqs = [
        ("Python Fundamentals", "Variables, functions, classes, modules, packages, and basic data structures. You should be comfortable writing Python code."),
        ("Basic SQL", "SELECT, INSERT, UPDATE, and JOIN operations. You don't need to be an expert, but understanding relational databases helps."),
        ("REST API Concepts", "Understanding of HTTP methods (GET, POST, PUT, DELETE), JSON format, and how web APIs work."),
        ("Git Basics", "Clone, commit, push, pull operations. We assume you can work with a Git repository."),
        ("Command Line", "Basic terminal/command prompt navigation. You'll need to run Python scripts and manage virtual environments.")
    ]

    for title, desc in prereqs:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # What we explain from scratch
    add_heading(doc, "What You Don't Need to Know (We Explain Everything)", 2)

    doc.add_paragraph(
        "The following topics are explained from first principles in this tutorial. "
        "You don't need any prior knowledge of:"
    )

    explained = [
        "What AI agents are and how they work",
        "Large Language Models (LLMs) and how they generate text",
        "Prompt engineering techniques",
        "LangChain and LangGraph frameworks",
        "Vector embeddings and similarity search",
        "FAISS and vector databases",
        "Retrieval-Augmented Generation (RAG)",
        "Multi-agent orchestration patterns",
        "FastAPI (beyond basic Python web knowledge)",
        "Docker containerization",
        "Cloud deployment"
    ]
    for item in explained:
        doc.add_paragraph(item, style='List Bullet')

    # How to read
    add_heading(doc, "How to Read This Tutorial", 2)

    p = doc.add_paragraph()
    run = p.add_run("Sequential Path (Recommended for Beginners): ")
    run.bold = True
    p.add_run(
        "Read the entire tutorial from start to finish. Each chapter builds on previous ones, "
        "and we introduce concepts in a carefully designed order. This path takes longer but "
        "gives you the deepest understanding."
    )

    p = doc.add_paragraph()
    run = p.add_run("Reference Path (For Experienced Developers): ")
    run.bold = True
    p.add_run(
        "If you're already familiar with AI concepts, you can jump to specific sections. "
        "Use the table of contents to find what you need. Each major section is relatively "
        "self-contained."
    )

    doc.add_paragraph("Suggested reading order based on your background:")

    headers = ["Your Background", "Start Here", "Focus On"]
    rows = [
        ["New to AI/ML", "Chapter 1", "Parts 1-3 (Fundamentals)"],
        ["Know ML, new to LLMs", "Chapter 2", "Parts 2-4 (Implementation)"],
        ["Know LLMs, new to agents", "Chapter 3", "Part 3 (Methods Deep Dive)"],
        ["Know agents, need deployment", "Chapter 16", "Parts 4-5 (Production)"],
    ]
    create_table(doc, headers, rows)

    doc.add_page_break()

def create_table_of_contents(doc):
    """Create detailed table of contents"""
    add_heading(doc, "Table of Contents", 1)

    toc = [
        ("PART 1: FOUNDATIONS OF AI AGENTS", [
            ("Chapter 1: Introduction to AI Agents", [
                "1.1 What is an AI Agent?",
                "1.2 A Simple Mental Model for Agents",
                "1.3 Why Agents Matter in Modern Software",
                "1.4 Types of AI Agents",
                "1.5 Building Your First Simple Agent (Code Example)"
            ]),
            ("Chapter 2: Understanding Large Language Models", [
                "2.1 What is an LLM?",
                "2.2 How LLMs Generate Text (Step by Step)",
                "2.3 The Transformer Architecture (Simplified)",
                "2.4 Key Parameters: Temperature, Tokens, Context",
                "2.5 LLM Capabilities and Limitations",
                "2.6 Choosing the Right LLM for Your Application"
            ]),
            ("Chapter 3: Prompt Engineering Fundamentals", [
                "3.1 What is Prompt Engineering?",
                "3.2 The Anatomy of a Good Prompt",
                "3.3 Prompt Patterns That Work",
                "3.4 Common Mistakes and How to Avoid Them",
                "3.5 Testing and Iterating on Prompts"
            ]),
            ("Chapter 4: Multi-Agent Systems", [
                "4.1 Why Multiple Agents?",
                "4.2 Agent Communication Patterns",
                "4.3 State Management Across Agents",
                "4.4 Orchestration Strategies",
                "4.5 Error Handling in Multi-Agent Systems"
            ]),
            ("Chapter 5: Retrieval-Augmented Generation (RAG)", [
                "5.1 The Knowledge Problem in LLMs",
                "5.2 What is RAG?",
                "5.3 Vector Embeddings Explained",
                "5.4 Similarity Search Fundamentals",
                "5.5 Building a RAG Pipeline from Scratch"
            ])
        ]),
        ("PART 2: THE CUSTOMER SUPPORT PROBLEM", [
            ("Chapter 6: Problem Definition and Requirements", [
                "6.1 The Business Case for AI Support",
                "6.2 Functional Requirements",
                "6.3 Non-Functional Requirements",
                "6.4 Success Metrics"
            ]),
            ("Chapter 7: System Design", [
                "7.1 High-Level Architecture",
                "7.2 Data Flow Design",
                "7.3 Component Interactions",
                "7.4 Technology Choices"
            ]),
            ("Chapter 8: Data Architecture", [
                "8.1 Knowledge Base Design",
                "8.2 Conversation Data Model",
                "8.3 User and Feedback Models",
                "8.4 Analytics Data Structure"
            ])
        ]),
        ("PART 3: BUILDING THE AGENTS", [
            ("Chapter 9: Query Categorization Agent", [
                "9.1 Purpose and Design",
                "9.2 The Categorization Prompt",
                "9.3 Implementation Details",
                "9.4 Testing and Edge Cases"
            ]),
            ("Chapter 10: Sentiment Analysis Agent", [
                "10.1 Understanding Sentiment in Support",
                "10.2 The Sentiment Prompt",
                "10.3 Priority Score Calculation",
                "10.4 Implementation Details"
            ]),
            ("Chapter 11: Knowledge Base Retrieval Agent", [
                "11.1 Vector Store Implementation",
                "11.2 The Retrieval Process",
                "11.3 Filtering and Ranking",
                "11.4 Performance Optimization"
            ]),
            ("Chapter 12: Response Agents", [
                "12.1 Technical Support Agent",
                "12.2 Billing Support Agent",
                "12.3 Account Support Agent",
                "12.4 General Support Agent"
            ]),
            ("Chapter 13: Escalation Agent", [
                "13.1 When to Escalate",
                "13.2 Escalation Logic",
                "13.3 Human Handoff Process",
                "13.4 Preserving Context"
            ]),
            ("Chapter 14: Workflow Orchestration with LangGraph", [
                "14.1 Introduction to LangGraph",
                "14.2 Building the Workflow",
                "14.3 Conditional Routing",
                "14.4 Error Handling"
            ])
        ]),
        ("PART 4: PRODUCTION IMPLEMENTATION", [
            ("Chapter 15: Project Structure", [
                "15.1 Directory Organization",
                "15.2 Module Dependencies",
                "15.3 Configuration Management"
            ]),
            ("Chapter 16: Database Implementation", [
                "16.1 SQLAlchemy Models",
                "16.2 Database Connections",
                "16.3 Query Functions",
                "16.4 Migrations"
            ]),
            ("Chapter 17: API Development", [
                "17.1 FastAPI Setup",
                "17.2 Request/Response Schemas",
                "17.3 Endpoint Implementation",
                "17.4 Error Handling"
            ]),
            ("Chapter 18: Webhook System", [
                "18.1 Webhook Architecture",
                "18.2 Event Types",
                "18.3 Delivery with Retry",
                "18.4 Security (HMAC Signatures)"
            ]),
            ("Chapter 19: Testing Strategy", [
                "19.1 Test Organization",
                "19.2 Unit Tests",
                "19.3 Integration Tests",
                "19.4 Running Tests"
            ]),
            ("Chapter 20: Deployment", [
                "20.1 Docker Configuration",
                "20.2 Environment Variables",
                "20.3 Railway Deployment",
                "20.4 Monitoring and Logging"
            ])
        ]),
        ("PART 5: RESULTS AND ANALYSIS", [
            ("Chapter 21: Performance Metrics", [
                "21.1 Response Time Analysis",
                "21.2 Accuracy Metrics",
                "21.3 System Statistics"
            ]),
            ("Chapter 22: Lessons Learned", [
                "22.1 What Worked Well",
                "22.2 Challenges Encountered",
                "22.3 Future Improvements"
            ])
        ]),
        ("APPENDICES", [
            ("Appendix A: Complete Code Listings", []),
            ("Appendix B: API Reference", []),
            ("Appendix C: Glossary of Terms", []),
            ("Appendix D: Troubleshooting Guide", []),
            ("Appendix E: Further Reading", [])
        ])
    ]

    for part_name, chapters in toc:
        p = doc.add_paragraph()
        run = p.add_run(part_name)
        run.bold = True
        run.font.size = Pt(12)

        for chapter_name, sections in chapters:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(chapter_name)
            run.bold = True

            for section in sections:
                p = doc.add_paragraph(section)
                p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

# ============================================================================
# PART 2: FOUNDATIONS OF AI AGENTS
# ============================================================================

def create_chapter_1(doc):
    """Chapter 1: Introduction to AI Agents"""
    add_heading(doc, "PART 1: FOUNDATIONS OF AI AGENTS", 1)
    doc.add_paragraph()
    add_heading(doc, "Chapter 1: Introduction to AI Agents", 1)

    add_section_intro(doc,
        "Before we dive into building a complex multi-agent system, we need to understand "
        "what an AI agent actually is. This chapter provides the foundational knowledge "
        "you'll need throughout the rest of this tutorial."
    )

    # 1.1 What is an AI Agent?
    add_heading(doc, "1.1 What is an AI Agent?", 2)

    doc.add_paragraph(
        "An AI agent is a software component that can perceive its environment, make decisions, "
        "and take actions to achieve specific goals. Unlike traditional software that follows "
        "predetermined rules, an agent uses artificial intelligence to determine what to do "
        "based on the situation it encounters."
    )

    doc.add_paragraph(
        "Let's break this down with a concrete example. Consider a customer support scenario:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Traditional Software Approach:")
    run.bold = True

    add_code_block(doc, '''# Traditional rule-based approach
def handle_query(query):
    if "password" in query.lower():
        return "To reset your password, go to Settings > Security"
    elif "refund" in query.lower():
        return "Refunds are processed within 5-7 business days"
    else:
        return "Please contact support for assistance"''')

    doc.add_paragraph(
        "This approach has severe limitations. It can only handle exact keyword matches, "
        "can't understand context or nuance, and requires manual rules for every possible scenario."
    )

    p = doc.add_paragraph()
    run = p.add_run("AI Agent Approach:")
    run.bold = True

    add_code_block(doc, '''# AI agent approach
def handle_query(query):
    # Agent perceives: understands the query using NLP
    understanding = analyze_query(query)

    # Agent decides: determines best course of action
    if understanding.needs_escalation:
        action = "escalate_to_human"
    elif understanding.category == "technical":
        action = "provide_technical_help"
    else:
        action = "provide_general_help"

    # Agent acts: generates appropriate response
    response = generate_response(query, action, context)
    return response''')

    doc.add_paragraph(
        "The AI agent can understand queries it has never seen before, adapt its response "
        "based on context, and handle the infinite variety of ways customers express themselves."
    )

    add_key_insight(doc,
        "An AI agent is defined by three capabilities: perception (understanding input), "
        "decision-making (choosing what to do), and action (producing output). Traditional "
        "software typically only has the action part, with decisions hardcoded by developers."
    )

    # 1.2 A Simple Mental Model for Agents
    add_heading(doc, "1.2 A Simple Mental Model for Agents", 2)

    doc.add_paragraph(
        "To understand agents intuitively, let's use some analogies from everyday life and "
        "software development:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Analogy 1: A Human Customer Service Representative")
    run.bold = True

    doc.add_paragraph(
        "Think about what a human support rep does: They listen to the customer (perception), "
        "think about how to help based on their training and knowledge (decision), and then "
        "respond appropriately (action). An AI agent does the same thing, but using algorithms "
        "instead of human cognition."
    )

    p = doc.add_paragraph()
    run = p.add_run("Analogy 2: A Software Function with Superpowers")
    run.bold = True

    doc.add_paragraph(
        "From a developer's perspective, you can think of an agent as a function that:"
    )

    bullets = [
        "Takes unstructured input (like natural language) instead of just structured data",
        "Can handle inputs it wasn't explicitly programmed for",
        "Uses a language model as its 'brain' to reason about what to do",
        "Produces contextually appropriate output rather than templated responses"
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run("Analogy 3: Microservices Architecture")
    run.bold = True

    doc.add_paragraph(
        "If you're familiar with microservices, agents are similar in philosophy. Just as "
        "microservices break a monolithic application into specialized services that communicate "
        "via APIs, a multi-agent system breaks AI processing into specialized agents that "
        "communicate via shared state."
    )

    headers = ["Microservices", "Multi-Agent System"]
    rows = [
        ["Service", "Agent"],
        ["API calls", "State passing"],
        ["Service registry", "Orchestrator/Workflow"],
        ["Each service has one job", "Each agent has one specialty"],
        ["Services are independently deployable", "Agents are independently testable"]
    ]
    create_table(doc, headers, rows)

    # 1.3 Why Agents Matter
    add_heading(doc, "1.3 Why Agents Matter in Modern Software", 2)

    doc.add_paragraph(
        "AI agents represent a fundamental shift in how we build software. Here's why they matter:"
    )

    p = doc.add_paragraph()
    run = p.add_run("1. Handling Unstructured Input")
    run.bold = True

    doc.add_paragraph(
        "Traditional software requires structured input - forms, APIs with defined schemas, "
        "button clicks. But humans naturally communicate in unstructured ways - natural language, "
        "images, voice. Agents can bridge this gap, accepting human communication and translating "
        "it into structured actions."
    )

    p = doc.add_paragraph()
    run = p.add_run("2. Adaptability Without Reprogramming")
    run.bold = True

    doc.add_paragraph(
        "When business rules change, traditional software requires code changes, testing, and "
        "deployment. Agents can often adapt through prompt changes or knowledge base updates, "
        "without touching code. Need to handle a new type of query? Update the prompt. "
        "Have new policies? Add them to the knowledge base."
    )

    p = doc.add_paragraph()
    run = p.add_run("3. Scaling Human Capabilities")
    run.bold = True

    doc.add_paragraph(
        "Some tasks traditionally required human judgment - understanding customer intent, "
        "generating helpful responses, knowing when to escalate. Agents can perform these tasks "
        "at scale, handling thousands of interactions simultaneously while maintaining quality."
    )

    p = doc.add_paragraph()
    run = p.add_run("4. Composability")
    run.bold = True

    doc.add_paragraph(
        "Agents can be composed together to handle complex workflows. A categorization agent "
        "feeds into a sentiment agent, which feeds into a response agent. This modularity "
        "makes systems easier to build, test, and maintain."
    )

    # 1.4 Types of AI Agents
    add_heading(doc, "1.4 Types of AI Agents", 2)

    doc.add_paragraph(
        "There are several types of AI agents, each suited for different tasks:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Simple Reflex Agents")
    run.bold = True

    doc.add_paragraph(
        "These agents respond to the current input without considering history. They're fast "
        "but limited. Example: A spam filter that classifies each email independently."
    )

    p = doc.add_paragraph()
    run = p.add_run("Model-Based Agents")
    run.bold = True

    doc.add_paragraph(
        "These maintain an internal model of the world and use it to make decisions. "
        "Example: A chatbot that remembers conversation history."
    )

    p = doc.add_paragraph()
    run = p.add_run("Goal-Based Agents")
    run.bold = True

    doc.add_paragraph(
        "These work toward specific goals, planning actions to achieve them. "
        "Example: A task automation agent that breaks down a goal into steps."
    )

    p = doc.add_paragraph()
    run = p.add_run("Utility-Based Agents")
    run.bold = True

    doc.add_paragraph(
        "These optimize for a utility function, choosing actions that maximize expected value. "
        "Example: A recommendation agent that maximizes user engagement."
    )

    p = doc.add_paragraph()
    run = p.add_run("Learning Agents")
    run.bold = True

    doc.add_paragraph(
        "These improve their performance over time through experience. "
        "Example: A support agent that learns from user feedback."
    )

    add_note(doc,
        "SmartSupport AI uses Model-Based Agents - each agent considers the conversation "
        "context and previous agent outputs when making decisions."
    )

    # 1.5 Building Your First Simple Agent
    add_heading(doc, "1.5 Building Your First Simple Agent (Code Example)", 2)

    doc.add_paragraph(
        "Let's build a simple AI agent from scratch to solidify these concepts. This agent "
        "will classify text into categories - a simplified version of our categorization agent."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 1: Install Dependencies")
    run.bold = True

    add_code_block(doc, '''pip install langchain-groq python-dotenv''')

    p = doc.add_paragraph()
    run = p.add_run("Step 2: Set Up Your Environment")
    run.bold = True

    add_code_block(doc, '''# Create a .env file with your API key:
GROQ_API_KEY=your_api_key_here''')

    p = doc.add_paragraph()
    run = p.add_run("Step 3: Create the Agent")
    run.bold = True

    add_code_block(doc, '''"""
simple_agent.py - Your first AI agent
"""
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()

# Initialize the LLM
llm = ChatGroq(
    temperature=0,  # Deterministic output
    groq_api_key=os.getenv("GROQ_API_KEY"),
    model_name="llama-3.3-70b-versatile"
)

# Define the agent's prompt
CLASSIFICATION_PROMPT = ChatPromptTemplate.from_template(
    """You are a text classification agent.

    Classify the following text into one of these categories:
    - Question: The text is asking something
    - Statement: The text is stating a fact or opinion
    - Command: The text is requesting an action
    - Greeting: The text is a greeting or farewell

    Text: {text}

    Respond with only the category name.
    Category:"""
)

# Create the processing chain
chain = CLASSIFICATION_PROMPT | llm | StrOutputParser()

def classify_text(text: str) -> str:
    """
    Our simple agent function.

    Perception: Receives text input
    Decision: Uses LLM to determine category
    Action: Returns the classification
    """
    result = chain.invoke({"text": text})
    return result.strip()

# Test the agent
if __name__ == "__main__":
    test_cases = [
        "How do I reset my password?",
        "The app crashed yesterday.",
        "Please send me a refund.",
        "Hello, good morning!"
    ]

    for text in test_cases:
        category = classify_text(text)
        print(f"Text: {text}")
        print(f"Category: {category}")
        print()''')

    p = doc.add_paragraph()
    run = p.add_run("Step 4: Run and Test")
    run.bold = True

    add_code_block(doc, '''python simple_agent.py

# Expected output:
# Text: How do I reset my password?
# Category: Question
#
# Text: The app crashed yesterday.
# Category: Statement
#
# Text: Please send me a refund.
# Category: Command
#
# Text: Hello, good morning!
# Category: Greeting''')

    doc.add_paragraph(
        "Congratulations! You've just built your first AI agent. Let's analyze what makes this an agent:"
    )

    bullets = [
        "Perception: It receives unstructured text input that it has never seen before",
        "Decision: It uses an LLM to reason about what category the text belongs to",
        "Action: It returns a classification that downstream code can use"
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    add_key_insight(doc,
        "This simple agent demonstrates the core pattern: prompt template + LLM + output parsing. "
        "Every agent in SmartSupport AI follows this same pattern, just with more sophisticated "
        "prompts and additional context."
    )

    doc.add_page_break()

def create_chapter_2(doc):
    """Chapter 2: Understanding Large Language Models"""
    add_heading(doc, "Chapter 2: Understanding Large Language Models", 1)

    add_section_intro(doc,
        "Large Language Models (LLMs) are the 'brains' of our AI agents. Understanding how they "
        "work - at least at a conceptual level - is essential for building effective agents. "
        "This chapter demystifies LLMs without requiring a machine learning background."
    )

    # 2.1 What is an LLM?
    add_heading(doc, "2.1 What is an LLM?", 2)

    doc.add_paragraph(
        "A Large Language Model is a type of artificial intelligence that has been trained on "
        "enormous amounts of text data - books, websites, code, conversations - to understand "
        "and generate human language. The key characteristics are:"
    )

    p = doc.add_paragraph()
    run = p.add_run('"Large": ')
    run.bold = True
    p.add_run(
        "These models have billions of parameters (learned values). GPT-3 has 175 billion "
        "parameters, Llama 3 has up to 70 billion. These parameters encode patterns and "
        "knowledge learned from training data."
    )

    p = doc.add_paragraph()
    run = p.add_run('"Language": ')
    run.bold = True
    p.add_run(
        "They're specifically designed for text - understanding it, generating it, and "
        "reasoning about it. They work with human languages (English, Spanish, etc.) "
        "as well as programming languages."
    )

    p = doc.add_paragraph()
    run = p.add_run('"Model": ')
    run.bold = True
    p.add_run(
        "In machine learning, a 'model' is a mathematical function that transforms inputs "
        "into outputs. An LLM transforms input text into output text based on patterns "
        "it learned during training."
    )

    doc.add_paragraph(
        "To give you a sense of scale, here's how much data these models are trained on:"
    )

    headers = ["Model", "Parameters", "Training Data", "Training Cost (Est.)"]
    rows = [
        ["GPT-3", "175 billion", "~500 billion tokens", "$4.6 million"],
        ["Llama 2 70B", "70 billion", "2 trillion tokens", "$2+ million"],
        ["Llama 3 70B", "70 billion", "15+ trillion tokens", "Undisclosed"],
    ]
    create_table(doc, headers, rows)

    add_note(doc,
        "SmartSupport AI uses Llama 3.3-70B via Groq's API. We get the benefits of a "
        "state-of-the-art model without the cost of training or hosting it ourselves."
    )

    # 2.2 How LLMs Generate Text
    add_heading(doc, "2.2 How LLMs Generate Text (Step by Step)", 2)

    doc.add_paragraph(
        "Understanding the text generation process helps you write better prompts and "
        "debug issues. Here's what happens when you send a query to an LLM:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 1: Tokenization")
    run.bold = True

    doc.add_paragraph(
        "The input text is broken into 'tokens' - typically words or parts of words. "
        "For example:"
    )

    add_code_block(doc, '''Input: "How do I reset my password?"
Tokens: ["How", " do", " I", " reset", " my", " password", "?"]

# Note: Some words might be split into sub-tokens:
Input: "unbelievable"
Tokens: ["un", "believ", "able"]''')

    doc.add_paragraph(
        "Tokenization matters because LLMs have a maximum context length measured in tokens "
        "(not characters or words). Llama 3 supports up to 128,000 tokens."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 2: Embedding")
    run.bold = True

    doc.add_paragraph(
        "Each token is converted into a vector (a list of numbers) that represents its "
        "meaning in context. These are called embeddings."
    )

    add_code_block(doc, '''# Conceptually (actual vectors are much larger):
"password" → [0.23, -0.15, 0.87, 0.42, ...]  # 4096 numbers
"reset" → [0.31, 0.08, 0.65, -0.22, ...]     # Similar dimension''')

    doc.add_paragraph(
        "Similar words have similar vectors. 'password' and 'credentials' would have "
        "vectors that are close together in this high-dimensional space."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 3: Attention Mechanism")
    run.bold = True

    doc.add_paragraph(
        "This is the key innovation of transformer models. The attention mechanism allows "
        "the model to consider relationships between all parts of the input simultaneously. "
        "When processing 'reset my password', the model can attend to the relationship "
        "between 'reset' and 'password' to understand the context."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 4: Token-by-Token Generation")
    run.bold = True

    doc.add_paragraph(
        "The model generates output one token at a time. For each position, it predicts "
        "the probability distribution over all possible tokens, selects one, and then "
        "uses that to predict the next token."
    )

    add_code_block(doc, '''Input: "How do I reset my password?"

Generation process:
Position 1: P("To") = 0.35, P("You") = 0.25, P("First") = 0.15, ...
            Selected: "To"

Position 2: P(" reset") = 0.40, P(" change") = 0.20, ...
            Selected: " reset"

Position 3: P(" your") = 0.45, P(" the") = 0.25, ...
            Selected: " your"

... continues until complete response ...

Final output: "To reset your password, go to Settings > Security..."''')

    add_key_insight(doc,
        "LLMs don't 'understand' in the human sense - they're sophisticated pattern matchers. "
        "They've seen so many examples of password reset instructions that they can generate "
        "plausible ones, but they don't 'know' what a password actually is."
    )

    # 2.3 Key Parameters
    add_heading(doc, "2.3 Key Parameters: Temperature, Tokens, Context", 2)

    doc.add_paragraph(
        "When using LLMs, several parameters significantly affect their behavior:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Temperature (0.0 - 2.0)")
    run.bold = True

    doc.add_paragraph(
        "Controls randomness in token selection. Lower = more deterministic, higher = more creative."
    )

    add_code_block(doc, '''# Temperature examples for "The sky is ___":

temperature=0.0 (deterministic):
  Always picks highest probability: "blue"

temperature=0.7 (balanced):
  Might pick: "blue", "clear", "cloudy", "beautiful"

temperature=1.5 (creative):
  Might pick: "azure", "endless", "whispering", "painted"''')

    doc.add_paragraph(
        "For customer support, we use temperature=0.0 because we want consistent, "
        "reliable responses. The same question should get the same answer."
    )

    p = doc.add_paragraph()
    run = p.add_run("Max Tokens")
    run.bold = True

    doc.add_paragraph(
        "The maximum number of tokens in the generated response. This prevents runaway "
        "generation and controls costs (you pay per token with most APIs)."
    )

    add_code_block(doc, '''# In SmartSupport AI configuration:
llm_max_tokens: int = 1000  # Enough for detailed response

# Too low (50 tokens): Response gets cut off mid-sentence
# Too high (10000 tokens): Wastes money, might generate unnecessary content''')

    p = doc.add_paragraph()
    run = p.add_run("Context Window")
    run.bold = True

    doc.add_paragraph(
        "The maximum total tokens (input + output) the model can handle at once. "
        "This limits how much conversation history and context you can include."
    )

    headers = ["Model", "Context Window", "Practical Limit"]
    rows = [
        ["GPT-3.5", "4,096 tokens", "~3,000 words"],
        ["GPT-4", "8,192 tokens", "~6,000 words"],
        ["GPT-4 Turbo", "128,000 tokens", "~96,000 words"],
        ["Llama 3 70B", "128,000 tokens", "~96,000 words"],
    ]
    create_table(doc, headers, rows)

    # 2.4 LLM Capabilities and Limitations
    add_heading(doc, "2.4 LLM Capabilities and Limitations", 2)

    doc.add_paragraph(
        "Understanding what LLMs can and cannot do is crucial for building reliable systems."
    )

    p = doc.add_paragraph()
    run = p.add_run("What LLMs Do Well:")
    run.bold = True

    capabilities = [
        ("Text Understanding", "Grasping meaning, intent, and nuance in natural language"),
        ("Text Generation", "Producing coherent, contextually appropriate text"),
        ("Classification", "Categorizing text into predefined classes"),
        ("Summarization", "Condensing long text while preserving key information"),
        ("Translation", "Converting between languages"),
        ("Code Generation", "Writing and explaining code"),
        ("Following Instructions", "Performing tasks described in natural language")
    ]

    for cap, desc in capabilities:
        p = doc.add_paragraph()
        run = p.add_run(f"{cap}: ")
        run.bold = True
        p.add_run(desc)

    p = doc.add_paragraph()
    run = p.add_run("What LLMs Struggle With:")
    run.bold = True

    limitations = [
        ("Mathematical Reasoning", "Complex calculations, especially multi-step math"),
        ("Factual Accuracy", "They can 'hallucinate' - generate plausible but false information"),
        ("Real-Time Information", "Knowledge is frozen at training time"),
        ("Consistency", "May give different answers to the same question with high temperature"),
        ("Long-Term Planning", "Struggle with complex, multi-step reasoning"),
        ("Self-Awareness", "Can't truly know when they don't know something")
    ]

    for lim, desc in limitations:
        p = doc.add_paragraph()
        run = p.add_run(f"{lim}: ")
        run.bold = True
        p.add_run(desc)

    add_warning(doc,
        "Hallucination is the biggest risk in customer support. An LLM might confidently "
        "describe a refund policy that doesn't exist. This is why we use RAG - to ground "
        "responses in actual documentation."
    )

    # 2.5 Choosing the Right LLM
    add_heading(doc, "2.5 Choosing the Right LLM for Your Application", 2)

    doc.add_paragraph(
        "For SmartSupport AI, we chose Groq's Llama 3.3-70B. Here's our decision process:"
    )

    headers = ["Factor", "Our Requirement", "Llama 3.3-70B via Groq"]
    rows = [
        ["Quality", "Near GPT-4 level", "Comparable to GPT-3.5/4"],
        ["Speed", "< 1 second responses", "~100 tokens/sec (very fast)"],
        ["Cost", "Affordable for startup", "Free tier available"],
        ["Privacy", "Data not used for training", "No training on queries"],
        ["Availability", "High uptime", "99.9% SLA"],
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph(
        "Alternative options we considered:"
    )

    alternatives = [
        ("OpenAI GPT-4", "Higher quality but more expensive, slower, vendor lock-in"),
        ("Anthropic Claude", "Excellent for long context, but less available"),
        ("Local Llama", "Full control but requires GPU infrastructure"),
        ("Fine-tuned model", "Could be better for our domain but requires training data and expertise")
    ]

    for name, reason in alternatives:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(reason)

    doc.add_page_break()

def create_chapter_3(doc):
    """Chapter 3: Prompt Engineering Fundamentals"""
    add_heading(doc, "Chapter 3: Prompt Engineering Fundamentals", 1)

    add_section_intro(doc,
        "Prompt engineering is the art and science of crafting instructions that get "
        "the desired behavior from an LLM. It's one of the most important skills for "
        "building effective AI agents. This chapter teaches you the fundamentals."
    )

    # 3.1 What is Prompt Engineering?
    add_heading(doc, "3.1 What is Prompt Engineering?", 2)

    doc.add_paragraph(
        "A 'prompt' is the text you send to an LLM. 'Prompt engineering' is the process "
        "of designing prompts that reliably produce the outputs you want. Think of it as "
        "writing very precise instructions for a highly capable but literal assistant."
    )

    doc.add_paragraph(
        "Here's a simple example showing how prompt design affects output:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Poor Prompt:")
    run.bold = True

    add_code_block(doc, '''Prompt: "Categorize this: My app keeps crashing"
Output: "This appears to be a technical issue related to application
        stability. The user is experiencing repeated crashes which
        could be caused by..."

Problem: Too verbose, not easily parseable by code''')

    p = doc.add_paragraph()
    run = p.add_run("Good Prompt:")
    run.bold = True

    add_code_block(doc, '''Prompt: "Categorize this customer query into exactly one category:
        Technical, Billing, Account, or General.

        Query: My app keeps crashing

        Respond with only the category name, nothing else.
        Category:"

Output: "Technical"

Result: Clean, parseable output''')

    doc.add_paragraph(
        "The difference is dramatic. The second prompt:"
    )

    bullets = [
        "Specifies the exact categories to choose from",
        "Makes clear that only ONE category should be chosen",
        "Explicitly requests only the category name",
        "Ends with 'Category:' to prime the model for a short response"
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # 3.2 The Anatomy of a Good Prompt
    add_heading(doc, "3.2 The Anatomy of a Good Prompt", 2)

    doc.add_paragraph(
        "Every effective prompt has several components. Let's examine a real prompt from "
        "SmartSupport AI:"
    )

    add_code_block(doc, '''"""You are an expert customer support query classifier.

Categorize the following customer query into ONE of these categories:
- Technical: Issues with software, hardware, bugs, errors, setup, configuration
- Billing: Payment issues, invoices, refunds, subscriptions, pricing
- Account: Login, password, profile, account settings, security
- General: Company policies, general inquiries, feedback

Query: {query}

{context}

Respond with ONLY the category name (Technical, Billing, Account, or General).
Category:"""''')

    doc.add_paragraph("Let's break down each component:")

    p = doc.add_paragraph()
    run = p.add_run("1. Role Definition (Line 1)")
    run.bold = True

    doc.add_paragraph(
        '"You are an expert customer support query classifier" - This establishes the '
        "persona the LLM should adopt. Using 'expert' tends to produce more confident, "
        "accurate responses. The specific role focuses the model on the task at hand."
    )

    p = doc.add_paragraph()
    run = p.add_run("2. Task Description (Lines 3-7)")
    run.bold = True

    doc.add_paragraph(
        '"Categorize the following customer query into ONE of these categories" - '
        "Clear instruction of what to do. The emphasis on 'ONE' prevents the model "
        "from hedging with multiple categories."
    )

    p = doc.add_paragraph()
    run = p.add_run("3. Category Definitions with Examples (Lines 4-7)")
    run.bold = True

    doc.add_paragraph(
        "Each category includes examples of what belongs there. This reduces ambiguity. "
        "Without examples, 'Account' vs 'Technical' would be unclear for login issues."
    )

    p = doc.add_paragraph()
    run = p.add_run("4. Input Placeholder (Line 9)")
    run.bold = True

    doc.add_paragraph(
        '"{query}" - This is where the actual user input goes. Using placeholders '
        "makes the prompt reusable as a template."
    )

    p = doc.add_paragraph()
    run = p.add_run("5. Context Placeholder (Line 11)")
    run.bold = True

    doc.add_paragraph(
        '"{context}" - Optional additional context like conversation history. '
        "This helps with queries like 'Still not working' that reference previous messages."
    )

    p = doc.add_paragraph()
    run = p.add_run("6. Output Format Specification (Line 13)")
    run.bold = True

    doc.add_paragraph(
        '"Respond with ONLY the category name" - Explicit instruction about output format. '
        "This prevents verbose explanations and makes parsing reliable."
    )

    p = doc.add_paragraph()
    run = p.add_run('7. Output Priming (Line 14)')
    run.bold = True

    doc.add_paragraph(
        '"Category:" - Ending with this primes the model to complete with just the category. '
        "It's like filling in a blank."
    )

    # 3.3 Prompt Patterns That Work
    add_heading(doc, "3.3 Prompt Patterns That Work", 2)

    doc.add_paragraph(
        "Over time, practitioners have discovered patterns that consistently improve results:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Pattern 1: Role-Task-Format (RTF)")
    run.bold = True

    add_code_block(doc, '''# Structure:
# 1. Define the role
# 2. Describe the task
# 3. Specify the output format

"""You are a [ROLE].

[TASK DESCRIPTION]

[OUTPUT FORMAT SPECIFICATION]"""

# Example:
"""You are a sentiment analysis expert.

Analyze the emotional tone of the following text and classify it as
Positive, Neutral, Negative, or Angry.

Text: {text}

Respond with only the sentiment label.
Sentiment:"""''')

    p = doc.add_paragraph()
    run = p.add_run("Pattern 2: Few-Shot Examples")
    run.bold = True

    add_code_block(doc, '''# Provide examples of desired input-output pairs

"""Classify the customer intent:

Example 1:
Input: "How much does the premium plan cost?"
Intent: Pricing Inquiry

Example 2:
Input: "I can't log into my account"
Intent: Technical Issue

Example 3:
Input: "I want to cancel my subscription"
Intent: Cancellation Request

Now classify this:
Input: {query}
Intent:"""''')

    p = doc.add_paragraph()
    run = p.add_run("Pattern 3: Chain of Thought")
    run.bold = True

    add_code_block(doc, '''# Ask the model to reason step by step

"""Determine if this customer query should be escalated to a human agent.

Query: {query}
Sentiment: {sentiment}
Priority: {priority}

Think through this step by step:
1. Is the sentiment angry or very negative?
2. Is the priority score 8 or higher?
3. Does the query mention legal action, demands, or explicit escalation?
4. Has the customer made multiple attempts without resolution?

Based on your reasoning, should this be escalated? (Yes/No)
Escalate:"""''')

    p = doc.add_paragraph()
    run = p.add_run("Pattern 4: Constraints and Boundaries")
    run.bold = True

    add_code_block(doc, '''# Explicitly state what NOT to do

"""Generate a response to this billing inquiry.

Query: {query}

Guidelines:
- Keep response under 200 words
- Do NOT promise specific refund amounts
- Do NOT share internal policy details
- If unsure, direct to human support
- Use empathetic language for frustrated customers

Response:"""''')

    # 3.4 Common Mistakes
    add_heading(doc, "3.4 Common Mistakes and How to Avoid Them", 2)

    p = doc.add_paragraph()
    run = p.add_run("Mistake 1: Ambiguous Instructions")
    run.bold = True

    add_code_block(doc, '''# Bad - ambiguous
"Categorize this message appropriately"

# Good - specific
"Categorize this message into exactly one of these categories:
Technical, Billing, Account, General"''')

    p = doc.add_paragraph()
    run = p.add_run("Mistake 2: No Output Format Specification")
    run.bold = True

    add_code_block(doc, '''# Bad - no format specified
"What category is this query?"
# Might return: "This query appears to be related to technical issues..."

# Good - format specified
"What category is this query? Respond with only the category name."
# Returns: "Technical"''')

    p = doc.add_paragraph()
    run = p.add_run("Mistake 3: Too Much in One Prompt")
    run.bold = True

    add_code_block(doc, '''# Bad - trying to do everything at once
"Categorize this query, analyze sentiment, check if it needs
escalation, and generate a response."

# Good - one task per prompt (use multiple agents)
Agent 1: "Categorize this query..."
Agent 2: "Analyze the sentiment..."
Agent 3: "Generate a response..."''')

    p = doc.add_paragraph()
    run = p.add_run("Mistake 4: Not Handling Edge Cases")
    run.bold = True

    add_code_block(doc, '''# Bad - no handling for unclear queries
"Categorize this: asdfghjkl"
# Model might hallucinate a category

# Good - explicit handling
"If the query is unclear, nonsensical, or doesn't fit any category,
respond with 'Unclear'. Otherwise, categorize as..."''')

    add_key_insight(doc,
        "The best prompts are unambiguous, specify exact output format, and handle edge cases. "
        "Treat prompt writing like writing a contract - assume nothing and specify everything."
    )

    doc.add_page_break()

def create_chapter_4(doc):
    """Chapter 4: Multi-Agent Systems"""
    add_heading(doc, "Chapter 4: Multi-Agent Systems", 1)

    add_section_intro(doc,
        "Now that you understand individual agents, let's explore how multiple agents "
        "work together. Multi-agent systems are more powerful, maintainable, and robust "
        "than single monolithic agents."
    )

    # 4.1 Why Multiple Agents?
    add_heading(doc, "4.1 Why Multiple Agents?", 2)

    doc.add_paragraph(
        "You might wonder why we don't just use one big agent that does everything. "
        "Here's a comparison:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Single Agent Approach:")
    run.bold = True

    add_code_block(doc, '''MEGA_PROMPT = """You are a customer support agent.

For the following query:
1. Determine the category (Technical/Billing/Account/General)
2. Analyze the sentiment (Positive/Neutral/Negative/Angry)
3. Calculate priority (1-10)
4. Search knowledge base for relevant articles
5. Generate an appropriate response
6. Decide if escalation is needed

Query: {query}

Provide all of the above in your response."""''')

    doc.add_paragraph("Problems with this approach:")

    problems = [
        "Prompt is complex and hard to maintain",
        "All tasks are coupled - can't update one without affecting others",
        "Can't test individual components",
        "If one part fails, everything fails",
        "Hard to debug which part went wrong",
        "Can't optimize individual steps"
    ]
    for p_text in problems:
        doc.add_paragraph(p_text, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run("Multi-Agent Approach:")
    run.bold = True

    add_code_block(doc, '''# Each agent has a focused, simple prompt

CATEGORIZATION_PROMPT = "Categorize into Technical/Billing/Account/General..."
SENTIMENT_PROMPT = "Analyze sentiment as Positive/Neutral/Negative/Angry..."
RESPONSE_PROMPT = "Generate helpful response based on category and sentiment..."

# Agents are orchestrated in a pipeline
result = (
    categorize(query)
    |> analyze_sentiment
    |> retrieve_kb
    |> generate_response
)''')

    doc.add_paragraph("Benefits of multi-agent approach:")

    benefits = [
        "Each prompt is simple and focused",
        "Can update categorization without touching response generation",
        "Each agent can be tested independently",
        "Failures are isolated and recoverable",
        "Easy to debug - check each agent's output",
        "Can optimize each step (e.g., use smaller model for categorization)"
    ]
    for b in benefits:
        doc.add_paragraph(b, style='List Bullet')

    # 4.2 Agent Communication Patterns
    add_heading(doc, "4.2 Agent Communication Patterns", 2)

    doc.add_paragraph(
        "How do agents share information? There are several patterns:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Pattern 1: Sequential Pipeline (What We Use)")
    run.bold = True

    add_code_block(doc, '''# Each agent passes state to the next
Query → Categorizer → Sentiment → KB Retrieval → Response → Output
        (state)       (state)     (state)        (state)

# State accumulates as it passes through:
Initial: {query: "..."}
After Categorizer: {query: "...", category: "Technical"}
After Sentiment: {query: "...", category: "Technical", sentiment: "Negative"}
After KB: {..., kb_results: [...]}
After Response: {..., response: "..."}''')

    p = doc.add_paragraph()
    run = p.add_run("Pattern 2: Parallel Execution")
    run.bold = True

    add_code_block(doc, '''# Independent agents run simultaneously
Query → [Categorizer, Sentiment] (parallel) → Merge → Response
                    ↓
        Both run at same time, results merged

# Faster but requires careful coordination''')

    p = doc.add_paragraph()
    run = p.add_run("Pattern 3: Hierarchical (Orchestrator Pattern)")
    run.bold = True

    add_code_block(doc, '''# Master agent delegates to specialists
                Orchestrator
               /     |     \\
         Technical  Billing  General
            Agent    Agent    Agent

# Orchestrator decides which specialist to invoke''')

    p = doc.add_paragraph()
    run = p.add_run("Pattern 4: Collaborative (Debate Pattern)")
    run.bold = True

    add_code_block(doc, '''# Multiple agents discuss and reach consensus
Agent A: "I think this is a billing issue"
Agent B: "The mention of 'login' suggests account issue"
Agent C: "Considering context, I agree with Agent B"
Final: Account issue (majority consensus)

# More robust but slower and more expensive''')

    doc.add_paragraph(
        "SmartSupport AI uses the Sequential Pipeline pattern because:"
    )

    reasons = [
        "Simple to understand and debug",
        "Natural fit for our workflow (categorize → analyze → respond)",
        "State naturally flows forward",
        "Easy to add new agents in the pipeline",
        "LangGraph provides excellent support for this pattern"
    ]
    for r in reasons:
        doc.add_paragraph(r, style='List Bullet')

    # 4.3 State Management
    add_heading(doc, "4.3 State Management Across Agents", 2)

    doc.add_paragraph(
        "The key challenge in multi-agent systems is managing shared state. "
        "Here's how we do it:"
    )

    add_code_block(doc, '''# From src/agents/state.py
from typing import TypedDict, Optional, List, Dict, Any

class AgentState(TypedDict):
    """
    Shared state passed between all agents.
    Each agent reads what it needs and adds its contribution.
    """

    # === INPUT (set at start) ===
    query: str                  # The customer's query
    user_id: str                # Who is asking
    conversation_id: str        # Unique conversation identifier

    # === ANALYSIS (set by analysis agents) ===
    category: Optional[str]     # Set by Categorizer
    sentiment: Optional[str]    # Set by Sentiment Analyzer
    priority_score: Optional[int]  # Set by Sentiment Analyzer

    # === CONTEXT (for enriching responses) ===
    user_context: Optional[Dict[str, Any]]  # User history, VIP status
    conversation_history: Optional[List[Dict[str, str]]]  # Previous messages
    kb_results: Optional[List[Dict[str, Any]]]  # Retrieved FAQs

    # === OUTPUT (set by response agents) ===
    response: Optional[str]     # The generated response

    # === ROUTING (set by escalation check) ===
    should_escalate: bool       # Whether to route to human
    escalation_reason: Optional[str]  # Why escalating

    # === METADATA (for analytics) ===
    metadata: Optional[Dict[str, Any]]
    processing_time: Optional[float]''')

    doc.add_paragraph(
        "This design has several advantages:"
    )

    advantages = [
        "Type safety: TypedDict provides IDE autocompletion and type checking",
        "Clear ownership: Each field is documented with which agent sets it",
        "Immutable pattern: Agents return new state rather than mutating",
        "Easy debugging: Can inspect state at any point in pipeline",
        "Extensible: Add new fields without breaking existing agents"
    ]
    for a in advantages:
        doc.add_paragraph(a, style='List Bullet')

    add_code_block(doc, '''# How agents use state:

def categorize_query(state: AgentState) -> AgentState:
    """Categorizer reads query, writes category"""
    query = state["query"]  # Read

    # ... LLM call to categorize ...

    state["category"] = "Technical"  # Write
    return state

def analyze_sentiment(state: AgentState) -> AgentState:
    """Sentiment analyzer reads query+category, writes sentiment"""
    query = state["query"]  # Read
    category = state["category"]  # Read (from previous agent)

    # ... LLM call to analyze ...

    state["sentiment"] = "Negative"  # Write
    state["priority_score"] = 6      # Write
    return state''')

    # 4.4 Error Handling
    add_heading(doc, "4.4 Error Handling in Multi-Agent Systems", 2)

    doc.add_paragraph(
        "What happens when an agent fails? We need graceful degradation:"
    )

    add_code_block(doc, '''def categorize_query(state: AgentState) -> AgentState:
    try:
        # Normal processing
        llm_manager = get_llm_manager()
        raw_category = llm_manager.invoke_with_retry(
            CATEGORIZATION_PROMPT,
            {"query": state["query"]}
        )
        category = parse_llm_category(raw_category)
        state["category"] = category

    except Exception as e:
        # Graceful fallback
        app_logger.error(f"Categorization failed: {e}")
        state["category"] = "General"  # Safe default

        # Optionally flag for review
        if not state.get("metadata"):
            state["metadata"] = {}
        state["metadata"]["categorization_failed"] = True

    return state  # Always return state, never raise''')

    doc.add_paragraph("Our error handling strategy:")

    strategies = [
        ("Safe Defaults", "Each agent has a fallback value (e.g., 'General' for category)"),
        ("Continue Pipeline", "Failures don't stop the pipeline - next agent gets state"),
        ("Flag Issues", "Metadata tracks which agents had problems"),
        ("Retry Logic", "LLM calls include automatic retry with backoff"),
        ("Logging", "All errors are logged for debugging")
    ]

    for name, desc in strategies:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    add_key_insight(doc,
        "In production, graceful degradation is essential. A customer getting a 'General' "
        "response when categorization fails is far better than an error message or no response at all."
    )

    doc.add_page_break()

def create_chapter_5(doc):
    """Chapter 5: Retrieval-Augmented Generation (RAG)"""
    add_heading(doc, "Chapter 5: Retrieval-Augmented Generation (RAG)", 1)

    add_section_intro(doc,
        "RAG is one of the most important techniques for building reliable AI systems. "
        "It solves the 'knowledge problem' - how to give LLMs accurate, up-to-date, "
        "company-specific information."
    )

    # 5.1 The Knowledge Problem
    add_heading(doc, "5.1 The Knowledge Problem in LLMs", 2)

    doc.add_paragraph(
        "LLMs have a fundamental limitation: their knowledge is frozen at training time. "
        "This creates several problems for customer support:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Problem 1: Outdated Information")
    run.bold = True

    doc.add_paragraph(
        "An LLM trained in January 2024 knows nothing about features you released in "
        "March 2024. If a customer asks about your new feature, the LLM will either "
        "say it doesn't know or (worse) make something up."
    )

    p = doc.add_paragraph()
    run = p.add_run("Problem 2: No Company-Specific Knowledge")
    run.bold = True

    doc.add_paragraph(
        "LLMs know general facts about the world but nothing about YOUR specific products, "
        "pricing, policies, or procedures. Asking 'How do I enable 2FA?' might get generic "
        "instructions that don't match your actual app."
    )

    p = doc.add_paragraph()
    run = p.add_run("Problem 3: Hallucination")
    run.bold = True

    doc.add_paragraph(
        "When LLMs don't know something, they often generate plausible-sounding but "
        "completely false information. This is called 'hallucination'. An LLM might "
        "confidently describe a refund policy that doesn't exist."
    )

    add_code_block(doc, '''# Example of hallucination risk:

Customer: "What's your refund policy?"

# Without RAG (LLM makes up a policy):
Response: "We offer a 45-day money-back guarantee on all purchases..."
Reality: Your actual policy is 30 days

# With RAG (LLM uses your actual documentation):
Retrieved: "Refund policy: 30-day money-back guarantee for new subscribers"
Response: "We offer a 30-day money-back guarantee for new subscribers..."''')

    add_warning(doc,
        "Hallucination in customer support can lead to legal issues, customer frustration, "
        "and loss of trust. RAG is not optional for production systems - it's essential."
    )

    # 5.2 What is RAG?
    add_heading(doc, "5.2 What is RAG?", 2)

    doc.add_paragraph(
        "Retrieval-Augmented Generation (RAG) is a technique that combines information "
        "retrieval with text generation. Instead of relying solely on what the LLM "
        "learned during training, we:"
    )

    doc.add_paragraph("1. Store company knowledge in a searchable database", style='List Number')
    doc.add_paragraph("2. When a query comes in, search for relevant information", style='List Number')
    doc.add_paragraph("3. Include the retrieved information in the LLM's prompt", style='List Number')
    doc.add_paragraph("4. LLM generates a response grounded in the retrieved facts", style='List Number')

    add_code_block(doc, '''# The RAG Pipeline Visualized:

Customer Query: "How do I enable two-factor authentication?"
                            │
                            ▼
            ┌───────────────────────────────┐
            │  1. ENCODE QUERY TO VECTOR    │
            │  "How do I enable 2FA?" →     │
            │  [0.23, -0.15, 0.87, ...]     │
            └───────────────────────────────┘
                            │
                            ▼
            ┌───────────────────────────────┐
            │  2. SEARCH VECTOR DATABASE    │
            │  Find similar FAQ vectors     │
            │  Returns top 3 matches        │
            └───────────────────────────────┘
                            │
                            ▼
            ┌───────────────────────────────┐
            │  3. RETRIEVED DOCUMENTS       │
            │  FAQ #8: "How to enable 2FA"  │
            │  Score: 0.92 (very relevant)  │
            └───────────────────────────────┘
                            │
                            ▼
            ┌───────────────────────────────┐
            │  4. AUGMENT LLM PROMPT        │
            │  "Based on this info: [FAQ]   │
            │   Answer: How do I enable..." │
            └───────────────────────────────┘
                            │
                            ▼
            ┌───────────────────────────────┐
            │  5. GENERATE RESPONSE         │
            │  LLM creates answer using     │
            │  the retrieved information    │
            └───────────────────────────────┘
                            │
                            ▼
            Response: "To enable 2FA in our app:
                      1. Go to Settings > Security
                      2. Click 'Two-Factor Authentication'
                      ..." (accurate, from your docs)''')

    # 5.3 Vector Embeddings
    add_heading(doc, "5.3 Vector Embeddings Explained", 2)

    doc.add_paragraph(
        "The magic of RAG lies in 'vector embeddings' - a way to represent text as "
        "numbers that capture meaning."
    )

    p = doc.add_paragraph()
    run = p.add_run("What is an Embedding?")
    run.bold = True

    doc.add_paragraph(
        "An embedding is a list of numbers (a 'vector') that represents the meaning of text. "
        "Texts with similar meanings have similar embeddings (vectors that are 'close' together)."
    )

    p = doc.add_paragraph()
    run = p.add_run("Analogy: GPS Coordinates for Meaning")
    run.bold = True

    doc.add_paragraph(
        "Think of embeddings like GPS coordinates, but for meaning instead of location:"
    )

    add_code_block(doc, '''# Physical locations have coordinates:
Paris: (48.8566° N, 2.3522° E)
London: (51.5074° N, 0.1278° W)
New York: (40.7128° N, 74.0060° W)

# Nearby cities have similar coordinates
# Paris and London are closer to each other than to New York

# Similarly, text meanings have "coordinates" (embeddings):
"happy": [0.82, 0.15, -0.34, 0.67, ...]  # 384 numbers
"joyful": [0.79, 0.18, -0.31, 0.65, ...]  # Very similar!
"sad": [-0.45, 0.22, 0.78, -0.33, ...]    # Very different

# "happy" and "joyful" are close in meaning-space
# "happy" and "sad" are far apart''')

    p = doc.add_paragraph()
    run = p.add_run("How Embeddings Are Created:")
    run.bold = True

    doc.add_paragraph(
        "We use a model called a 'sentence transformer' that was trained specifically "
        "to produce meaningful embeddings:"
    )

    add_code_block(doc, '''from sentence_transformers import SentenceTransformer

# Load the embedding model
encoder = SentenceTransformer("all-MiniLM-L6-v2")

# Generate embeddings
texts = [
    "How do I reset my password?",
    "I forgot my login credentials",
    "What's your refund policy?"
]

embeddings = encoder.encode(texts)

# embeddings is now a numpy array of shape (3, 384)
# Each text is represented by 384 numbers

print(embeddings.shape)  # (3, 384)
print(embeddings[0][:5])  # First 5 numbers of first embedding
# Output: [ 0.023, -0.156,  0.087,  0.042, -0.231]''')

    p = doc.add_paragraph()
    run = p.add_run("Why 'all-MiniLM-L6-v2'?")
    run.bold = True

    doc.add_paragraph("We chose this model because:")

    reasons = [
        "Fast: Small model (22M parameters) runs quickly",
        "Good quality: Trained specifically for semantic similarity",
        "384 dimensions: Good balance of expressiveness and efficiency",
        "Open source: Free to use, no API costs",
        "Well-tested: Widely used in production systems"
    ]
    for r in reasons:
        doc.add_paragraph(r, style='List Bullet')

    # 5.4 Similarity Search
    add_heading(doc, "5.4 Similarity Search Fundamentals", 2)

    doc.add_paragraph(
        "Once we have embeddings, we need to find which stored documents are most "
        "similar to the user's query. This is called 'similarity search'."
    )

    p = doc.add_paragraph()
    run = p.add_run("Measuring Similarity:")
    run.bold = True

    doc.add_paragraph(
        "There are several ways to measure how similar two vectors are:"
    )

    add_code_block(doc, '''# 1. Euclidean Distance (L2)
# Measures straight-line distance between vectors
# Lower = more similar

import numpy as np

def euclidean_distance(a, b):
    return np.sqrt(np.sum((a - b) ** 2))

# 2. Cosine Similarity
# Measures the angle between vectors
# Higher = more similar (range: -1 to 1)

def cosine_similarity(a, b):
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))

# We use FAISS with L2 distance, then convert to similarity:
similarity = 1 / (1 + distance)
# distance=0 → similarity=1 (identical)
# distance=1 → similarity=0.5
# distance=9 → similarity=0.1''')

    p = doc.add_paragraph()
    run = p.add_run("FAISS: Fast Similarity Search")
    run.bold = True

    doc.add_paragraph(
        "FAISS (Facebook AI Similarity Search) is a library that makes similarity "
        "search fast, even with millions of vectors:"
    )

    add_code_block(doc, '''import faiss
import numpy as np

# Create an index for 384-dimensional vectors
dimension = 384
index = faiss.IndexFlatL2(dimension)

# Add document embeddings to the index
document_embeddings = encoder.encode(documents)  # Shape: (N, 384)
index.add(document_embeddings.astype('float32'))

# Search for similar documents
query = "How do I change my password?"
query_embedding = encoder.encode([query])  # Shape: (1, 384)

# Find top 3 most similar documents
k = 3
distances, indices = index.search(query_embedding.astype('float32'), k)

# distances: [[0.23, 0.45, 0.67]]  # Lower = more similar
# indices: [[5, 12, 3]]  # Document indices

# Get the actual documents
for i, (dist, idx) in enumerate(zip(distances[0], indices[0])):
    similarity = 1 / (1 + dist)
    print(f"Match {i+1}: Document {idx}, Similarity: {similarity:.2f}")
    print(f"  Content: {documents[idx][:100]}...")''')

    # 5.5 Building a RAG Pipeline
    add_heading(doc, "5.5 Building a RAG Pipeline from Scratch", 2)

    doc.add_paragraph(
        "Let's build a complete RAG pipeline step by step. This is a simplified version "
        "of what SmartSupport AI uses."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 1: Prepare Your Knowledge Base")
    run.bold = True

    add_code_block(doc, '''# knowledge_base.json
{
  "faqs": [
    {
      "id": 1,
      "question": "How do I reset my password?",
      "answer": "To reset your password: 1) Click 'Forgot Password'..."
    },
    {
      "id": 2,
      "question": "What is your refund policy?",
      "answer": "We offer a 30-day money-back guarantee..."
    }
  ]
}''')

    p = doc.add_paragraph()
    run = p.add_run("Step 2: Create the Vector Store")
    run.bold = True

    add_code_block(doc, '''import json
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer

class SimpleRAG:
    def __init__(self):
        # Load embedding model
        self.encoder = SentenceTransformer("all-MiniLM-L6-v2")
        self.dimension = 384

        # Initialize FAISS index
        self.index = faiss.IndexFlatL2(self.dimension)
        self.documents = []

    def load_knowledge_base(self, filepath):
        """Load and index FAQs"""
        with open(filepath) as f:
            data = json.load(f)

        for faq in data["faqs"]:
            # Combine question and answer for better matching
            text = f"Q: {faq['question']}\\nA: {faq['answer']}"
            self.documents.append({
                "id": faq["id"],
                "question": faq["question"],
                "answer": faq["answer"],
                "text": text
            })

        # Generate embeddings for all documents
        texts = [doc["text"] for doc in self.documents]
        embeddings = self.encoder.encode(texts)

        # Add to FAISS index
        self.index.add(embeddings.astype('float32'))

        print(f"Indexed {len(self.documents)} documents")''')

    p = doc.add_paragraph()
    run = p.add_run("Step 3: Implement Search")
    run.bold = True

    add_code_block(doc, '''    def search(self, query, k=3):
        """Search for relevant documents"""
        # Encode the query
        query_embedding = self.encoder.encode([query])

        # Search FAISS
        distances, indices = self.index.search(
            query_embedding.astype('float32'), k
        )

        # Build results with similarity scores
        results = []
        for dist, idx in zip(distances[0], indices[0]):
            if idx < len(self.documents):
                doc = self.documents[idx].copy()
                doc["similarity"] = float(1 / (1 + dist))
                results.append(doc)

        return results''')

    p = doc.add_paragraph()
    run = p.add_run("Step 4: Generate Response with Context")
    run.bold = True

    add_code_block(doc, '''    def generate_response(self, query, llm):
        """RAG: Retrieve then Generate"""
        # Step 1: Retrieve relevant documents
        results = self.search(query, k=3)

        # Step 2: Build context from retrieved docs
        context = "Relevant information from our knowledge base:\\n"
        for i, doc in enumerate(results):
            context += f"{i+1}. {doc['question']}: {doc['answer'][:200]}...\\n"

        # Step 3: Create augmented prompt
        prompt = f"""Use the following information to answer the question.

{context}

Question: {query}

Provide a helpful answer based on the information above. If the information
doesn't fully answer the question, say so and provide what help you can.

Answer:"""

        # Step 4: Generate response
        response = llm.invoke(prompt)
        return response, results''')

    p = doc.add_paragraph()
    run = p.add_run("Step 5: Use the Complete System")
    run.bold = True

    add_code_block(doc, '''# Initialize RAG system
rag = SimpleRAG()
rag.load_knowledge_base("knowledge_base.json")

# Initialize LLM
from langchain_groq import ChatGroq
llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0)

# Process a query
query = "How can I change my password?"
response, sources = rag.generate_response(query, llm)

print("Response:", response)
print("\\nSources used:")
for src in sources:
    print(f"  - {src['question']} (similarity: {src['similarity']:.2f})")''')

    add_key_insight(doc,
        "RAG transforms LLMs from 'knows everything poorly' to 'knows exactly what you "
        "tell it'. For customer support, this means accurate, up-to-date, company-specific "
        "responses that customers can trust."
    )

    doc.add_page_break()

# ============================================================================
# PART 2: THE CUSTOMER SUPPORT PROBLEM
# ============================================================================

def create_part_2(doc):
    """Create Part 2: The Customer Support Problem"""
    add_heading(doc, "PART 2: THE CUSTOMER SUPPORT PROBLEM", 1)
    doc.add_paragraph()

    # Chapter 6
    add_heading(doc, "Chapter 6: Problem Definition and Requirements", 1)

    add_section_intro(doc,
        "Now that we understand the foundational technologies, let's define the specific "
        "problem we're solving and what success looks like."
    )

    add_heading(doc, "6.1 The Business Case for AI Support", 2)

    doc.add_paragraph(
        "Customer support is expensive. Consider these typical costs:"
    )

    headers = ["Metric", "Traditional Support", "AI-Augmented Support"]
    rows = [
        ["Cost per interaction", "$5-15", "$0.01-0.10"],
        ["Availability", "Business hours", "24/7"],
        ["Response time", "Minutes to hours", "Seconds"],
        ["Scalability", "Linear (hire more)", "Near-infinite"],
        ["Consistency", "Variable", "Consistent"]
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph(
        "AI support doesn't replace humans - it handles the 70-80% of queries that are "
        "routine, freeing humans for complex issues requiring judgment and empathy."
    )

    add_heading(doc, "6.2 Functional Requirements", 2)

    doc.add_paragraph("SmartSupport AI must:")

    headers = ["ID", "Requirement", "Description", "Priority"]
    rows = [
        ["F1", "Accept queries", "Process natural language customer queries", "Must Have"],
        ["F2", "Categorize", "Classify into Technical/Billing/Account/General", "Must Have"],
        ["F3", "Analyze sentiment", "Detect Positive/Neutral/Negative/Angry", "Must Have"],
        ["F4", "Calculate priority", "Score 1-10 based on urgency", "Must Have"],
        ["F5", "Search knowledge", "Find relevant FAQs using semantic search", "Must Have"],
        ["F6", "Generate responses", "Create helpful, accurate answers", "Must Have"],
        ["F7", "Escalate", "Route complex issues to humans", "Must Have"],
        ["F8", "Track conversations", "Persist history in database", "Must Have"],
        ["F9", "Provide API", "REST endpoints for integration", "Must Have"],
        ["F10", "Web interface", "Browser-based chat UI", "Should Have"],
        ["F11", "Webhooks", "Notify external systems of events", "Should Have"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "6.3 Non-Functional Requirements", 2)

    headers = ["ID", "Requirement", "Target", "Rationale"]
    rows = [
        ["N1", "Response time", "< 2 seconds", "User expectation for chat"],
        ["N2", "Availability", "99.9%", "Support is critical"],
        ["N3", "Accuracy", "> 90%", "Correct categorization"],
        ["N4", "Escalation rate", "< 15%", "Most queries handled by AI"],
        ["N5", "Test coverage", "> 25%", "Code quality baseline"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "6.4 Success Metrics", 2)

    doc.add_paragraph("We measure success by:")

    metrics = [
        ("Response Time", "Average time from query to response"),
        ("Categorization Accuracy", "Percentage of correctly categorized queries"),
        ("Resolution Rate", "Queries resolved without human intervention"),
        ("Escalation Rate", "Queries requiring human handoff"),
        ("User Satisfaction", "Ratings from feedback system")
    ]

    for name, desc in metrics:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # Chapter 7
    add_heading(doc, "Chapter 7: System Design", 1)

    add_heading(doc, "7.1 High-Level Architecture", 2)

    add_code_block(doc, '''
┌─────────────────────────────────────────────────────────────────────────┐
│                          SMARTSUPPORT AI                                │
├─────────────────────────────────────────────────────────────────────────┤
│                                                                         │
│   ┌─────────┐     ┌─────────────────────────────────────────────────┐  │
│   │  User   │────▶│                   API Layer                      │  │
│   │Interface│◀────│              (FastAPI REST API)                  │  │
│   └─────────┘     └─────────────────────────────────────────────────┘  │
│                                      │                                  │
│                                      ▼                                  │
│   ┌─────────────────────────────────────────────────────────────────┐  │
│   │                    AGENT WORKFLOW (LangGraph)                    │  │
│   │                                                                   │  │
│   │  ┌────────────┐   ┌────────────┐   ┌────────────┐               │  │
│   │  │Categorizer │──▶│ Sentiment  │──▶│KB Retrieval│               │  │
│   │  └────────────┘   └────────────┘   └────────────┘               │  │
│   │                                           │                       │  │
│   │                                           ▼                       │  │
│   │  ┌────────────┐   ┌────────────────────────────────────────┐    │  │
│   │  │ Escalation │◀──│        Response Agents                  │    │  │
│   │  │   Check    │   │  Technical │ Billing │ Account │General │    │  │
│   │  └────────────┘   └────────────────────────────────────────┘    │  │
│   │                                                                   │  │
│   └─────────────────────────────────────────────────────────────────┘  │
│          │                    │                        │                │
│          ▼                    ▼                        ▼                │
│   ┌───────────┐       ┌─────────────┐         ┌─────────────┐         │
│   │   LLM     │       │  Vector DB  │         │  Database   │         │
│   │  (Groq)   │       │   (FAISS)   │         │(PostgreSQL) │         │
│   └───────────┘       └─────────────┘         └─────────────┘         │
│                                                                         │
└─────────────────────────────────────────────────────────────────────────┘
''')

    add_heading(doc, "7.2 Data Flow", 2)

    doc.add_paragraph("When a customer sends a query:")

    steps = [
        "Query arrives at API endpoint (/api/v1/query)",
        "API creates initial state with query, user_id, conversation_id",
        "Workflow starts at Categorizer agent",
        "Category is determined, state updated",
        "Sentiment analyzer processes query, calculates priority",
        "KB Retrieval searches for relevant FAQs",
        "Escalation check determines if human needed",
        "Appropriate response agent generates answer",
        "Response saved to database",
        "API returns response to user",
        "Webhooks triggered for external systems"
    ]

    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # Chapter 8
    add_heading(doc, "Chapter 8: Data Architecture", 1)

    add_heading(doc, "8.1 Knowledge Base Design", 2)

    doc.add_paragraph(
        "Our knowledge base contains 30 FAQs across 4 categories:"
    )

    headers = ["Category", "Count", "Example Topics"]
    rows = [
        ["Technical", "10", "App crashes, login issues, sync problems, performance"],
        ["Billing", "10", "Charges, refunds, subscriptions, pricing, payments"],
        ["Account", "5", "Password reset, profile updates, 2FA, deletion"],
        ["General", "5", "Business hours, contact info, security, platforms"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "8.2 Database Schema", 2)

    add_code_block(doc, '''# Core Tables:

┌──────────────────┐       ┌──────────────────┐
│      users       │       │  conversations   │
├──────────────────┤       ├──────────────────┤
│ id (PK)          │──┐    │ id (PK)          │
│ user_id (unique) │  │    │ conversation_id  │
│ name             │  │    │ user_id (FK)  ◀──┘
│ email            │  │    │ query            │
│ is_vip           │  │    │ category         │
│ created_at       │  │    │ sentiment        │
└──────────────────┘  │    │ priority_score   │
                      │    │ response         │
                      │    │ status           │
                      │    │ escalated        │
                      │    └──────────────────┘
                      │              │
                      │              ▼
                      │    ┌──────────────────┐
                      │    │    messages      │
                      │    ├──────────────────┤
                      │    │ id (PK)          │
                      │    │ conversation_id  │
                      │    │ role             │
                      │    │ content          │
                      │    │ created_at       │
                      │    └──────────────────┘
                      │
                      │    ┌──────────────────┐
                      │    │    feedback      │
                      └──▶ ├──────────────────┤
                           │ id (PK)          │
                           │ conversation_id  │
                           │ user_id (FK)     │
                           │ rating (1-5)     │
                           │ comment          │
                           └──────────────────┘''')

    doc.add_page_break()

# ============================================================================
# MAIN GENERATION FUNCTION
# ============================================================================

def generate_comprehensive_tutorial():
    """Generate the complete comprehensive tutorial"""
    print("Creating comprehensive tutorial document...")

    doc = Document()

    # Configure default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    # Generate each section
    print("  Creating title page...")
    create_title_page(doc)

    print("  Creating reader's guide...")
    create_readers_guide(doc)

    print("  Creating table of contents...")
    create_table_of_contents(doc)

    print("  Creating Chapter 1: Introduction to AI Agents...")
    create_chapter_1(doc)

    print("  Creating Chapter 2: Understanding LLMs...")
    create_chapter_2(doc)

    print("  Creating Chapter 3: Prompt Engineering...")
    create_chapter_3(doc)

    print("  Creating Chapter 4: Multi-Agent Systems...")
    create_chapter_4(doc)

    print("  Creating Chapter 5: RAG...")
    create_chapter_5(doc)

    print("  Creating Part 2: Problem Setup...")
    create_part_2(doc)

    # Save progress - Part 1 complete
    output_path = "tutorial_documentation/SmartSupport_AI_Comprehensive_Tutorial.docx"
    doc.save(output_path)
    print(f"\nPart 1 saved to: {output_path}")
    print("Run generate_part2() to continue...")

    return doc, output_path

if __name__ == "__main__":
    generate_comprehensive_tutorial()
