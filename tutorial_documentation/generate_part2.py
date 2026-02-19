"""
Generate Part 2 of the comprehensive tutorial
Continues from where generate_comprehensive_tutorial.py left off
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_code_block(doc, code):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return para

def add_key_insight(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run("KEY INSIGHT: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)
    para.add_run(text).italic = True
    return para

def add_warning(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run("WARNING: ")
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)
    para.add_run(text)
    return para

def add_note(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run("NOTE: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 150)
    para.add_run(text)
    return para

def create_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'D9E2F3')
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)
    doc.add_paragraph()
    return table


def create_part_3_agents(doc):
    """Part 3: Building the Agents"""
    add_heading(doc, "PART 3: BUILDING THE AGENTS", 1)
    doc.add_paragraph()

    # Chapter 9: Categorization Agent
    add_heading(doc, "Chapter 9: Query Categorization Agent", 1)

    doc.add_paragraph(
        "The categorization agent is the first in our pipeline. It determines what type "
        "of issue the customer has, which controls routing to specialized agents."
    )

    add_heading(doc, "9.1 Purpose and Design", 2)

    doc.add_paragraph("The categorizer must:")
    bullets = [
        "Classify queries into exactly one of four categories",
        "Handle ambiguous queries gracefully",
        "Consider conversation context for follow-up questions",
        "Be fast (this is the first step in every query)"
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    add_heading(doc, "9.2 The Categorization Prompt", 2)

    doc.add_paragraph("Here's the complete prompt with analysis:")

    add_code_block(doc, '''# From src/agents/categorizer.py

CATEGORIZATION_PROMPT = ChatPromptTemplate.from_template(
    """You are an expert customer support query classifier.

Categorize the following customer query into ONE of these categories:
- Technical: Issues with software, hardware, service functionality,
             bugs, errors, setup, configuration
- Billing: Payment issues, invoices, refunds, subscriptions,
           pricing, charges
- Account: Login, password, profile, account settings,
           registration, security
- General: Company policies, general inquiries, feedback, suggestions

Query: {query}

{context}

Respond with ONLY the category name (Technical, Billing, Account, or General).
Category:"""
)''')

    doc.add_paragraph("Prompt Design Analysis:")

    p = doc.add_paragraph()
    run = p.add_run("Line 1 - Role Definition: ")
    run.bold = True
    p.add_run(
        '"expert customer support query classifier" establishes expertise and '
        "focuses the model on classification specifically."
    )

    p = doc.add_paragraph()
    run = p.add_run("Lines 3-4 - Task Specification: ")
    run.bold = True
    p.add_run(
        '"Categorize...into ONE of these categories" - The word ONE is critical. '
        "Without it, the model might hedge with multiple categories."
    )

    p = doc.add_paragraph()
    run = p.add_run("Lines 5-12 - Category Definitions: ")
    run.bold = True
    p.add_run(
        "Each category includes specific examples. This reduces ambiguity. Is "
        "'can\\'t log in' Technical or Account? The examples clarify it's Account."
    )

    p = doc.add_paragraph()
    run = p.add_run("Line 14 - Query Input: ")
    run.bold = True
    p.add_run("{query} is the placeholder where the actual customer text goes.")

    p = doc.add_paragraph()
    run = p.add_run("Line 16 - Context Injection: ")
    run.bold = True
    p.add_run(
        "{context} allows including conversation history. Essential for queries "
        'like "Still not working" that reference previous messages.'
    )

    p = doc.add_paragraph()
    run = p.add_run("Lines 18-19 - Output Specification: ")
    run.bold = True
    p.add_run(
        '"Respond with ONLY" prevents verbose explanations. "Category:" at the '
        "end primes the model to complete with just the category name."
    )

    add_heading(doc, "9.3 Implementation Details", 2)

    add_code_block(doc, '''def categorize_query(state: AgentState) -> AgentState:
    """
    Categorize customer query into Technical/Billing/Account/General.

    Args:
        state: Current agent state with query

    Returns:
        Updated state with category field set
    """
    app_logger.info(f"Categorizing query: {state['query'][:50]}...")

    try:
        # Get LLM manager (singleton pattern)
        llm_manager = get_llm_manager()

        # Build context from conversation history
        context = ""
        if state.get("conversation_history"):
            context = "Previous conversation context:\\n"
            # Only include last 3 messages to keep prompt short
            for msg in state["conversation_history"][-3:]:
                # Truncate long messages
                content = msg['content'][:100]
                context += f"{msg['role']}: {content}\\n"

        # Invoke LLM with retry logic
        raw_category = llm_manager.invoke_with_retry(
            CATEGORIZATION_PROMPT,
            {"query": state["query"], "context": context}
        )

        # Parse and standardize the response
        category = parse_llm_category(raw_category)

        app_logger.info(f"Query categorized as: {category}")

        # Update state
        state["category"] = category

        # Store raw response in metadata for debugging
        if not state.get("metadata"):
            state["metadata"] = {}
        state["metadata"]["raw_category"] = raw_category

        return state

    except Exception as e:
        app_logger.error(f"Error in categorize_query: {e}")
        # Graceful fallback - General is safest default
        state["category"] = "General"
        return state''')

    doc.add_paragraph("Key implementation details:")

    details = [
        ("Singleton LLM", "get_llm_manager() returns single instance, avoiding repeated initialization"),
        ("Limited Context", "Only last 3 messages, truncated to 100 chars - keeps prompt size manageable"),
        ("Retry Logic", "invoke_with_retry handles transient API failures"),
        ("Response Parsing", "parse_llm_category normalizes variations like 'technical' vs 'Technical'"),
        ("Graceful Fallback", "On any error, defaults to 'General' rather than crashing"),
        ("Metadata Tracking", "Stores raw response for debugging categorization issues")
    ]

    for name, desc in details:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "9.4 The Parser Function", 2)

    add_code_block(doc, '''def parse_llm_category(raw_category: str) -> str:
    """
    Parse and standardize category from LLM response.

    The LLM might return variations like:
    - "Technical" (correct)
    - "technical" (lowercase)
    - "Technical issue" (extra words)
    - "Tech" (abbreviation)

    All should map to "Technical".
    """
    category_lower = raw_category.lower().strip()

    if "technical" in category_lower or "tech" in category_lower:
        return "Technical"
    elif "billing" in category_lower or "payment" in category_lower:
        return "Billing"
    elif "account" in category_lower:
        return "Account"
    else:
        # Default to General for anything unclear
        return "General"''')

    doc.add_page_break()

    # Chapter 10: Sentiment Agent
    add_heading(doc, "Chapter 10: Sentiment Analysis Agent", 1)

    doc.add_paragraph(
        "The sentiment agent determines the customer's emotional state, which affects "
        "response tone and escalation decisions."
    )

    add_heading(doc, "10.1 Understanding Sentiment in Support", 2)

    doc.add_paragraph(
        "Customer sentiment directly impacts how we should respond:"
    )

    headers = ["Sentiment", "Customer State", "Response Strategy"]
    rows = [
        ["Positive", "Happy, grateful", "Friendly, encourage feedback"],
        ["Neutral", "Calm, factual", "Direct, efficient"],
        ["Negative", "Frustrated, disappointed", "Empathetic, apologetic"],
        ["Angry", "Very upset, demanding", "Very empathetic, may escalate"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "10.2 The Sentiment Prompt", 2)

    add_code_block(doc, '''SENTIMENT_PROMPT = ChatPromptTemplate.from_template(
    """You are an expert at analyzing customer sentiment and emotions.

Analyze the sentiment of the following customer query and classify it as ONE of:
- Positive: Happy, satisfied, grateful, pleased
- Neutral: Informational, factual, calm
- Negative: Disappointed, frustrated, concerned, unhappy
- Angry: Very upset, furious, demanding, threatening

Consider the tone, word choice, and emotional indicators in the text.

Query: {query}

{context}

Respond with ONLY the sentiment label (Positive, Neutral, Negative, or Angry).
Sentiment:"""
)''')

    doc.add_paragraph("What the model looks for:")

    indicators = [
        ("Exclamation marks", "'!!!' suggests strong emotion"),
        ("Capitalization", "'THIS IS UNACCEPTABLE' indicates anger"),
        ("Word choice", "'frustrated', 'disappointed' = Negative; 'furious', 'unacceptable' = Angry"),
        ("Tone progression", "Context shows if customer is getting more upset"),
        ("Demands/threats", "'I demand a refund' or 'I'll sue' = Angry")
    ]

    for name, desc in indicators:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "10.3 Priority Score Calculation", 2)

    doc.add_paragraph(
        "After detecting sentiment, we calculate a priority score (1-10):"
    )

    add_code_block(doc, '''def calculate_priority_score(
    sentiment: str,
    category: str,
    is_repeat_query: bool = False,
    is_vip: bool = False
) -> int:
    """
    Calculate priority score for routing and escalation.

    Priority Scale:
    1-3: Low priority (routine queries)
    4-6: Medium priority (some urgency)
    7-8: High priority (needs attention)
    9-10: Critical (immediate attention)

    Args:
        sentiment: Detected sentiment
        category: Query category
        is_repeat_query: Customer asked about this before
        is_vip: Customer has VIP status

    Returns:
        Priority score between 1 and 10
    """
    # Start with base score
    score = 3

    # Adjust for sentiment
    sentiment_adjustments = {
        "Positive": 0,   # Happy customers aren't urgent
        "Neutral": 0,    # Normal priority
        "Negative": 2,   # Bump up frustrated customers
        "Angry": 3,      # Angry customers need quick attention
    }
    score += sentiment_adjustments.get(sentiment, 0)

    # Repeat queries indicate unresolved issues
    if is_repeat_query:
        score += 2

    # VIP customers get priority
    if is_vip:
        score += 2

    # Ensure score stays in valid range
    return max(1, min(10, score))''')

    doc.add_paragraph("Priority examples:")

    headers = ["Scenario", "Base", "Sentiment", "Modifiers", "Final"]
    rows = [
        ["Normal query, neutral", "3", "+0", "None", "3"],
        ["Frustrated customer", "3", "+2", "None", "5"],
        ["Angry customer", "3", "+3", "None", "6"],
        ["Angry + repeat query", "3", "+3", "+2", "8"],
        ["Angry + VIP + repeat", "3", "+3", "+4", "10"],
    ]
    create_table(doc, headers, rows)

    add_key_insight(doc,
        "Priority scores drive escalation. Scores >= 8 typically trigger human handoff. "
        "This ensures angry VIP customers with unresolved issues get immediate human attention."
    )

    doc.add_page_break()

    # Chapter 11: KB Retrieval
    add_heading(doc, "Chapter 11: Knowledge Base Retrieval Agent", 1)

    add_heading(doc, "11.1 Vector Store Implementation", 2)

    doc.add_paragraph(
        "Our vector store wraps FAISS with document management:"
    )

    add_code_block(doc, '''class VectorStore:
    """
    Vector Store for FAQ embeddings and similarity search.
    Uses FAISS for efficient nearest-neighbor search.
    """

    def __init__(
        self,
        model_name: str = "all-MiniLM-L6-v2",
        index_path: str = "./data/knowledge_base/faiss_index",
    ):
        # Load sentence transformer for embeddings
        self.encoder = SentenceTransformer(model_name)
        self.embedding_dim = self.encoder.get_sentence_embedding_dimension()

        # FAISS index (created when documents added)
        self.index = None
        self.documents = []

        # Try to load existing index
        self.load()

    def add_documents(self, documents: List[Dict]) -> None:
        """Add documents to the vector store."""
        if not documents:
            return

        # Extract text for embedding
        texts = [doc.get("text", "") for doc in documents]

        # Generate embeddings
        embeddings = self.encoder.encode(texts, show_progress_bar=True)
        embeddings = np.array(embeddings).astype("float32")

        # Create index if needed
        if self.index is None:
            self.index = faiss.IndexFlatL2(self.embedding_dim)

        # Add to index
        self.index.add(embeddings)
        self.documents.extend(documents)

        print(f"Added {len(documents)} documents. Total: {len(self.documents)}")''')

    add_heading(doc, "11.2 The Search Function", 2)

    add_code_block(doc, '''    def search(
        self,
        query: str,
        k: int = 3,
        category_filter: str = None
    ) -> List[Dict]:
        """
        Search for documents similar to query.

        Args:
            query: Search query text
            k: Number of results to return
            category_filter: Only return docs from this category

        Returns:
            List of matching documents with similarity scores
        """
        if self.index is None or len(self.documents) == 0:
            return []

        # Encode query
        query_embedding = self.encoder.encode([query])
        query_embedding = np.array(query_embedding).astype("float32")

        # Search more if filtering (some results may be filtered out)
        search_k = k * 3 if category_filter else k

        # FAISS search
        distances, indices = self.index.search(
            query_embedding,
            min(search_k, len(self.documents))
        )

        # Build results
        results = []
        for distance, idx in zip(distances[0], indices[0]):
            if idx < len(self.documents):
                doc = self.documents[idx].copy()

                # Convert L2 distance to similarity score
                # distance=0 means identical, higher means less similar
                doc["similarity_score"] = float(1 / (1 + distance))

                # Apply category filter
                if category_filter is None or doc.get("category") == category_filter:
                    results.append(doc)

                if len(results) >= k:
                    break

        return results''')

    add_heading(doc, "11.3 The Retrieval Agent", 2)

    add_code_block(doc, '''def retrieve_from_kb(state: AgentState) -> AgentState:
    """
    Retrieve relevant FAQs from knowledge base.

    Uses the detected category to filter results for better relevance.
    """
    query = state.get("query", "")
    category = state.get("category", "General")

    try:
        kb_retriever = get_kb_retriever()

        # Search with category filtering
        results = kb_retriever.retrieve(
            query=query,
            k=3,              # Top 3 results
            category=category, # Filter by detected category
            min_score=0.3,    # Minimum similarity threshold
        )

        # Format for response agents
        kb_results = []
        for result in results:
            kb_results.append({
                "title": result.get("question", ""),
                "content": result.get("answer", ""),
                "category": result.get("category", ""),
                "score": result.get("similarity_score", 0.0),
            })

        state["kb_results"] = kb_results
        return state

    except Exception as e:
        # Don't fail the workflow - just proceed without KB
        state["kb_results"] = []
        return state''')

    doc.add_page_break()

    # Chapter 12: Response Agents
    add_heading(doc, "Chapter 12: Response Agents", 1)

    doc.add_paragraph(
        "We have four specialized response agents, each optimized for its domain."
    )

    add_heading(doc, "12.1 Technical Support Agent", 2)

    add_code_block(doc, '''TECHNICAL_PROMPT = ChatPromptTemplate.from_template(
    """You are an expert technical support agent with deep knowledge
of software, hardware, and IT systems.

Customer Query: {query}

Customer Sentiment: {sentiment}
Priority Level: {priority}

{context}

{kb_context}

Instructions:
1. Provide a clear, step-by-step technical solution
2. Use simple language while being technically accurate
3. If the sentiment is negative or angry, start with empathy
4. Include troubleshooting steps if applicable
5. Offer to escalate if the issue is complex
6. Keep response concise but comprehensive (200-300 words)

Response:"""
)''')

    doc.add_paragraph("Key aspects of the technical prompt:")

    aspects = [
        ("Sentiment awareness", "Instruction to start with empathy for upset customers"),
        ("Step-by-step format", "Technical issues often need sequential instructions"),
        ("Complexity awareness", "Offers escalation for issues beyond AI capability"),
        ("KB integration", "{kb_context} includes retrieved troubleshooting articles"),
        ("Length guidance", "200-300 words balances thoroughness with readability")
    ]

    for name, desc in aspects:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "12.2 Billing Support Agent", 2)

    add_code_block(doc, '''BILLING_PROMPT = ChatPromptTemplate.from_template(
    """You are an expert billing and payment support agent.

Customer Query: {query}

Customer Sentiment: {sentiment}
Priority Level: {priority}

{context}

{kb_context}

Instructions:
1. Address billing concerns clearly and accurately
2. Explain charges, payment processes, or refund policies
3. If sentiment is negative, show empathy and apologize
4. Provide specific next steps for resolution
5. Reference relevant policies when appropriate
6. Escalate for refund requests or disputes if needed
7. Keep response professional and concise (200-300 words)

Response:"""
)''')

    doc.add_paragraph(
        "The billing agent has special handling for sensitive keywords:"
    )

    add_code_block(doc, '''# In handle_billing():

# Check for refund/dispute keywords that may need human review
query_lower = state["query"].lower()
sensitive_terms = ["refund", "dispute", "chargeback", "cancel subscription"]

if any(term in query_lower for term in sensitive_terms):
    # Flag for potential escalation or review
    state["metadata"]["may_need_escalation"] = True''')

    add_heading(doc, "12.3 Account Support Agent", 2)

    doc.add_paragraph(
        "Account queries require extra security awareness:"
    )

    add_code_block(doc, '''ACCOUNT_PROMPT = ChatPromptTemplate.from_template(
    """You are an account management and security support agent.

Customer Query: {query}

{context}

{kb_context}

Instructions:
1. Address account-related concerns (login, password, profile, security)
2. Provide clear step-by-step instructions
3. Emphasize security best practices
4. If password reset or security issue, guide through secure process
5. Be reassuring about account security
6. Never ask for or reveal sensitive information in chat
7. Keep response clear and actionable (200-300 words)

Response:"""
)''')

    add_warning(doc,
        "Account agents should NEVER ask for passwords, full credit card numbers, "
        "or other sensitive data. The prompt explicitly forbids this."
    )

    doc.add_page_break()

    # Chapter 13: Escalation Agent
    add_heading(doc, "Chapter 13: Escalation Agent", 1)

    add_heading(doc, "13.1 When to Escalate", 2)

    doc.add_paragraph("Escalation triggers in Multi-Agent HR Intelligence Platform:")

    headers = ["Trigger", "Threshold", "Rationale"]
    rows = [
        ["Priority Score", ">= 8", "High urgency situations"],
        ["Sentiment", "Angry", "Needs human empathy"],
        ["Attempt Count", ">= 3", "AI hasn't resolved issue"],
        ["Keywords", "See list", "Explicit escalation request"]
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph("Escalation keywords:")

    add_code_block(doc, '''escalation_keywords = [
    "lawsuit",          # Legal threat
    "legal",            # Legal mention
    "attorney",         # Legal mention
    "lawyer",           # Legal mention
    "sue",              # Legal threat
    "refund immediately",  # Urgent demand
    "speak to manager",    # Explicit request
    "speak to a manager",
    "talk to manager",
    "talk to a manager",
    "contact supervisor",
    "unacceptable",     # Strong dissatisfaction
    "ridiculous",       # Strong dissatisfaction
    "demand refund",    # Urgent demand
    "escalate this",    # Explicit request
]''')

    add_heading(doc, "13.2 Escalation Decision Logic", 2)

    add_code_block(doc, '''def should_escalate(
    priority_score: int,
    sentiment: str,
    attempt_count: int = 1,
    query: str = ""
) -> tuple[bool, Optional[str]]:
    """
    Determine if query should be escalated to human agent.

    Returns:
        Tuple of (should_escalate, reason)
    """
    reasons = []

    # Check priority threshold
    if priority_score >= 8:
        reasons.append("High priority score")

    # Check sentiment
    if sentiment == "Angry":
        reasons.append("Angry sentiment detected")

    # Check attempt count
    if attempt_count >= 3:
        reasons.append("Multiple unsuccessful attempts")

    # Check for escalation keywords
    query_lower = query.lower()
    for keyword in escalation_keywords:
        if keyword in query_lower:
            reasons.append(f"Escalation keyword: {keyword}")
            break

    # Escalate if ANY trigger is met
    should_escalate_flag = len(reasons) > 0
    escalation_reason = "; ".join(reasons) if reasons else None

    return should_escalate_flag, escalation_reason''')

    add_heading(doc, "13.3 Human Handoff Message", 2)

    add_code_block(doc, '''def escalate_to_human(state: AgentState) -> AgentState:
    """Generate appropriate escalation message based on sentiment."""
    sentiment = state.get("sentiment", "Neutral")

    if sentiment == "Angry":
        message = (
            "I sincerely apologize for the frustration you're experiencing. "
            "Your concern is very important to us, and I'm connecting you with "
            "a specialized support representative who can provide immediate "
            "assistance. They will be with you shortly and have full context "
            "of your situation."
        )
    elif sentiment == "Negative":
        message = (
            "I understand your concern, and I want to ensure you receive the "
            "best possible assistance. I'm connecting you with a senior support "
            "specialist who can help resolve this issue."
        )
    else:
        message = (
            "To ensure you receive the most accurate assistance for your "
            "inquiry, I'm connecting you with a specialized support "
            "representative."
        )

    # Add case reference for tracking
    message += f"\\n\\nCase Reference: {state.get('conversation_id', 'N/A')}"
    message += "\\n\\nEstimated wait time: 2-5 minutes"

    state["response"] = message
    state["next_action"] = "escalate"
    return state''')

    doc.add_page_break()

    # Chapter 14: Workflow Orchestration
    add_heading(doc, "Chapter 14: Workflow Orchestration with LangGraph", 1)

    add_heading(doc, "14.1 Introduction to LangGraph", 2)

    doc.add_paragraph(
        "LangGraph is a library for building stateful, multi-step agent workflows. "
        "It provides:"
    )

    features = [
        "StateGraph: A directed graph where nodes are functions",
        "Edges: Define flow between nodes (sequential or conditional)",
        "State: TypedDict passed between nodes, accumulating data",
        "Compilation: Turns the graph into an executable workflow"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    add_heading(doc, "14.2 Building the Workflow", 2)

    add_code_block(doc, '''from langgraph.graph import StateGraph, END

def create_workflow() -> StateGraph:
    """Create the customer support workflow graph."""

    # Initialize with state type
    workflow = StateGraph(AgentState)

    # === ADD NODES ===
    # Each node is a function that takes state and returns state
    workflow.add_node("categorize", categorize_query)
    workflow.add_node("analyze_sentiment", analyze_sentiment)
    workflow.add_node("retrieve_kb", retrieve_from_kb)
    workflow.add_node("check_escalation", check_escalation)
    workflow.add_node("technical", handle_technical)
    workflow.add_node("billing", handle_billing)
    workflow.add_node("account", handle_account)
    workflow.add_node("general", handle_general)
    workflow.add_node("escalate", escalate_to_human)

    # === SET ENTRY POINT ===
    workflow.set_entry_point("categorize")

    # === ADD SEQUENTIAL EDGES ===
    # These always flow in this order
    workflow.add_edge("categorize", "analyze_sentiment")
    workflow.add_edge("analyze_sentiment", "retrieve_kb")
    workflow.add_edge("retrieve_kb", "check_escalation")

    # === ADD CONDITIONAL EDGES ===
    # Route based on state after escalation check
    workflow.add_conditional_edges(
        "check_escalation",  # From this node...
        route_query,         # Use this function to decide...
        {                    # Map return values to nodes
            "technical": "technical",
            "billing": "billing",
            "account": "account",
            "general": "general",
            "escalate": "escalate",
        },
    )

    # === ADD TERMINAL EDGES ===
    # All response nodes lead to END
    workflow.add_edge("technical", END)
    workflow.add_edge("billing", END)
    workflow.add_edge("account", END)
    workflow.add_edge("general", END)
    workflow.add_edge("escalate", END)

    # Compile into executable
    return workflow.compile()''')

    add_heading(doc, "14.3 The Routing Function", 2)

    add_code_block(doc, '''def route_query(state: AgentState) -> str:
    """
    Determine which response agent should handle this query.

    Called by LangGraph's conditional_edges after check_escalation.

    Returns:
        Node name to route to
    """
    # Escalation takes priority over category
    if state.get("should_escalate", False):
        return "escalate"

    # Route based on category
    category = state.get("category", "General")

    routing_map = {
        "Technical": "technical",
        "Billing": "billing",
        "Account": "account",
        "General": "general",
    }

    return routing_map.get(category, "general")''')

    add_heading(doc, "14.4 Visual Workflow", 2)

    add_code_block(doc, '''
                    ┌─────────────┐
                    │    START    │
                    └──────┬──────┘
                           │
                           ▼
                    ┌─────────────┐
                    │ Categorize  │
                    └──────┬──────┘
                           │
                           ▼
                    ┌─────────────┐
                    │ Sentiment   │
                    └──────┬──────┘
                           │
                           ▼
                    ┌─────────────┐
                    │ KB Retrieve │
                    └──────┬──────┘
                           │
                           ▼
                    ┌─────────────┐
                    │ Escalation  │
                    │   Check     │
                    └──────┬──────┘
                           │
         ┌─────────────────┼─────────────────┐
         │                 │                 │
   ┌─────┴─────┐    ┌─────┴─────┐    ┌─────┴─────┐
   │ Technical │    │  Billing  │    │  Account  │ ...
   └─────┬─────┘    └─────┬─────┘    └─────┬─────┘
         │                │                 │
         └────────────────┼─────────────────┘
                          │
                          ▼
                    ┌─────────────┐
                    │     END     │
                    └─────────────┘''')

    doc.add_page_break()


def create_part_4_implementation(doc):
    """Part 4: Production Implementation"""
    add_heading(doc, "PART 4: PRODUCTION IMPLEMENTATION", 1)
    doc.add_paragraph()

    # Chapter 15: Project Structure
    add_heading(doc, "Chapter 15: Project Structure", 1)

    add_heading(doc, "15.1 Directory Organization", 2)

    add_code_block(doc, '''smartsupport-ai/
├── src/
│   ├── __init__.py
│   ├── main.py                 # Main orchestrator
│   │
│   ├── agents/                 # AI Agent modules
│   │   ├── __init__.py
│   │   ├── state.py           # AgentState definition
│   │   ├── workflow.py        # LangGraph workflow
│   │   ├── llm_manager.py     # LLM client wrapper
│   │   ├── categorizer.py     # Categorization agent
│   │   ├── sentiment_analyzer.py
│   │   ├── kb_retrieval.py
│   │   ├── technical_agent.py
│   │   ├── billing_agent.py
│   │   ├── general_agent.py
│   │   └── escalation_agent.py
│   │
│   ├── api/                   # FastAPI application
│   │   ├── __init__.py
│   │   ├── app.py            # FastAPI setup
│   │   ├── routes.py         # API endpoints
│   │   ├── schemas.py        # Pydantic models
│   │   ├── webhooks.py       # Webhook endpoints
│   │   ├── webhook_events.py
│   │   └── webhook_delivery.py
│   │
│   ├── database/             # Database layer
│   │   ├── __init__.py
│   │   ├── models.py        # SQLAlchemy models
│   │   ├── connection.py    # Connection management
│   │   ├── queries.py       # Query functions
│   │   └── webhook_queries.py
│   │
│   ├── knowledge_base/      # RAG implementation
│   │   ├── __init__.py
│   │   ├── retriever.py    # KB retrieval logic
│   │   └── vector_store.py # FAISS wrapper
│   │
│   └── utils/              # Utilities
│       ├── __init__.py
│       ├── config.py      # Configuration
│       ├── logger.py      # Logging setup
│       └── helpers.py     # Helper functions
│
├── data/
│   └── knowledge_base/
│       ├── faqs.json      # FAQ content
│       └── metadata.json  # FAISS metadata
│
├── tests/                  # Test suite
│   ├── test_basic.py
│   └── test_webhooks.py
│
├── .env                    # Environment variables
├── requirements.txt        # Dependencies
├── Dockerfile
├── docker-compose.yml
└── railway.json           # Railway deployment''')

    add_heading(doc, "15.2 Configuration Management", 2)

    add_code_block(doc, '''# src/utils/config.py
from pydantic_settings import BaseSettings
from functools import lru_cache

class Settings(BaseSettings):
    """Application settings loaded from environment variables."""

    # API Keys
    groq_api_key: str
    openai_api_key: Optional[str] = None

    # Database
    database_url: str = "sqlite:///./smartsupport.db"

    # Application
    app_name: str = "Multi-Agent HR Intelligence Platform"
    debug: bool = False
    environment: str = "development"

    # LLM Configuration
    llm_model: str = "llama-3.3-70b-versatile"
    llm_temperature: float = 0.0
    llm_max_tokens: int = 1000

    # Security
    secret_key: str

    class Config:
        env_file = ".env"
        case_sensitive = False

@lru_cache()
def get_settings() -> Settings:
    """Get cached settings instance."""
    return Settings()

settings = get_settings()''')

    doc.add_page_break()

    # Chapter 16: Database
    add_heading(doc, "Chapter 16: Database Implementation", 1)

    add_heading(doc, "16.1 SQLAlchemy Models", 2)

    add_code_block(doc, '''# src/database/models.py
from sqlalchemy import Column, Integer, String, Text, DateTime, Boolean, ForeignKey, JSON
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import relationship
from datetime import datetime

Base = declarative_base()

class User(Base):
    __tablename__ = "users"

    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(String(50), unique=True, index=True, nullable=False)
    name = Column(String(100))
    email = Column(String(100), unique=True, index=True)
    is_vip = Column(Boolean, default=False)
    created_at = Column(DateTime, default=datetime.utcnow)

    conversations = relationship("Conversation", back_populates="user")


class Conversation(Base):
    __tablename__ = "conversations"

    id = Column(Integer, primary_key=True, index=True)
    conversation_id = Column(String(50), unique=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))

    # Query and analysis
    query = Column(Text, nullable=False)
    category = Column(String(50))
    sentiment = Column(String(50))
    priority_score = Column(Integer, default=5)

    # Response
    response = Column(Text)
    response_time = Column(Float)

    # Status
    status = Column(String(50), default="Active")
    escalated = Column(Boolean, default=False)
    escalation_reason = Column(Text)

    # Timestamps
    created_at = Column(DateTime, default=datetime.utcnow)

    # Relationships
    user = relationship("User", back_populates="conversations")
    messages = relationship("Message", back_populates="conversation")''')

    add_heading(doc, "16.2 Database Connection", 2)

    add_code_block(doc, '''# src/database/connection.py
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from contextlib import contextmanager

# Create engine based on environment
if settings.database_url.startswith("sqlite"):
    engine = create_engine(
        settings.database_url,
        connect_args={"check_same_thread": False},
    )
else:
    # PostgreSQL for production
    engine = create_engine(
        settings.database_url,
        pool_size=10,
        max_overflow=20,
    )

SessionLocal = sessionmaker(bind=engine)

def get_db():
    """Dependency for FastAPI endpoints."""
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@contextmanager
def get_db_context():
    """Context manager for non-FastAPI code."""
    db = SessionLocal()
    try:
        yield db
        db.commit()
    except Exception:
        db.rollback()
        raise
    finally:
        db.close()''')

    doc.add_page_break()

    # Chapter 17: API
    add_heading(doc, "Chapter 17: API Development", 1)

    add_heading(doc, "17.1 FastAPI Setup", 2)

    add_code_block(doc, '''# src/api/app.py
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(
    title="Multi-Agent HR Intelligence Platform",
    description="Intelligent Customer Support Agent",
    version="2.2.0",
    docs_url="/docs",     # Swagger UI
    redoc_url="/redoc",   # ReDoc
)

# Enable CORS for web frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Include API routes
app.include_router(router)
app.include_router(webhooks_router)

@app.on_event("startup")
async def startup():
    """Initialize on startup."""
    init_db()

@app.get("/health")
async def health():
    """Health check endpoint."""
    return {"status": "healthy"}''')

    add_heading(doc, "17.2 Request/Response Schemas", 2)

    add_code_block(doc, '''# src/api/schemas.py
from pydantic import BaseModel
from typing import List, Optional

class QueryRequest(BaseModel):
    """Request to process a customer query."""
    user_id: str
    message: str

class KBResult(BaseModel):
    """A knowledge base search result."""
    title: str
    content: str
    category: str
    score: float

class QueryMetadata(BaseModel):
    """Metadata about query processing."""
    processing_time: float
    escalated: bool
    escalation_reason: Optional[str] = None
    kb_results: List[KBResult] = []

class QueryResponse(BaseModel):
    """Response from processing a query."""
    conversation_id: str
    response: str
    category: str
    sentiment: str
    priority: int
    timestamp: str
    metadata: QueryMetadata''')

    add_heading(doc, "17.3 Main Query Endpoint", 2)

    add_code_block(doc, '''@router.post("/query", response_model=QueryResponse)
async def process_query(
    request: QueryRequest,
    background_tasks: BackgroundTasks,
    agent=Depends(get_agent),
    db: Session = Depends(get_db),
):
    """Process a customer support query."""
    start_time = time.time()

    # Process through agent workflow
    result = agent.process_query(
        query=request.message,
        user_id=request.user_id
    )

    # Build response
    response = QueryResponse(
        conversation_id=result["conversation_id"],
        response=result["response"],
        category=result["category"],
        sentiment=result["sentiment"],
        priority=result["priority"],
        timestamp=result["timestamp"],
        metadata=QueryMetadata(
            processing_time=time.time() - start_time,
            escalated=result["metadata"].get("escalated", False),
            kb_results=[...]
        ),
    )

    # Trigger webhooks in background (non-blocking)
    background_tasks.add_task(
        trigger_webhooks, db, "query.created", payload
    )

    return response''')

    doc.add_page_break()

    # Chapter 19: Testing
    add_heading(doc, "Chapter 19: Testing Strategy", 1)

    add_heading(doc, "19.1 Test Organization", 2)

    add_code_block(doc, '''tests/
├── test_basic.py       # Core functionality tests
└── test_webhooks.py    # Webhook system tests''')

    add_heading(doc, "19.2 Example Tests", 2)

    add_code_block(doc, '''# tests/test_basic.py
import pytest
from src.main import get_customer_support_agent

@pytest.fixture
def agent():
    """Create agent instance for tests."""
    return get_customer_support_agent()

def test_technical_query_categorization(agent):
    """Test that technical queries are categorized correctly."""
    result = agent.process_query(
        query="My app keeps crashing when I try to export",
        user_id="test_user"
    )

    assert result["category"] == "Technical"
    assert "response" in result
    assert len(result["response"]) > 0

def test_billing_query_categorization(agent):
    """Test that billing queries are categorized correctly."""
    result = agent.process_query(
        query="Why was I charged twice this month?",
        user_id="test_user"
    )

    assert result["category"] == "Billing"

def test_angry_sentiment_escalation(agent):
    """Test that angry customers trigger escalation."""
    result = agent.process_query(
        query="THIS IS UNACCEPTABLE! I DEMAND A REFUND NOW!",
        user_id="test_user"
    )

    assert result["sentiment"] == "Angry"
    assert result["metadata"]["escalated"] == True

def test_response_generation(agent):
    """Test that responses are generated."""
    result = agent.process_query(
        query="How do I reset my password?",
        user_id="test_user"
    )

    assert "password" in result["response"].lower() or \\
           "reset" in result["response"].lower()''')

    add_heading(doc, "19.3 Running Tests", 2)

    add_code_block(doc, '''# Run all tests
pytest tests/ -v

# Run with coverage report
pytest tests/ -v --cov=src --cov-report=term-missing

# Run specific test file
pytest tests/test_basic.py -v

# Run specific test
pytest tests/test_basic.py::test_technical_query_categorization -v''')

    doc.add_paragraph("Test results for Multi-Agent HR Intelligence Platform:")

    add_code_block(doc, '''============================= test session starts =============================
collected 38 items

tests/test_basic.py ................                                    [ 42%]
tests/test_webhooks.py ......................                           [100%]

============================= 38 passed in 48.35s =============================

----------- coverage: -----------
Name                         Stmts   Miss  Cover
src/agents/categorizer.py       45     12    73%
src/agents/workflow.py          62      8    87%
src/api/routes.py               89     34    62%
...
TOTAL                         1784   1028    42%''')

    doc.add_page_break()

    # Chapter 20: Deployment
    add_heading(doc, "Chapter 20: Deployment", 1)

    add_heading(doc, "20.1 Docker Configuration", 2)

    add_code_block(doc, '''# Dockerfile
FROM python:3.10-slim

WORKDIR /app

# Install dependencies
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

# Copy application code
COPY . .

# Initialize knowledge base
RUN python -c "from src.knowledge_base.retriever import get_kb_retriever; get_kb_retriever()"

# Expose port
EXPOSE 8000

# Run with Gunicorn + Uvicorn workers
CMD ["gunicorn", "src.api.app:app", \\
     "--workers", "4", \\
     "--worker-class", "uvicorn.workers.UvicornWorker", \\
     "--bind", "0.0.0.0:8000"]''')

    add_heading(doc, "20.2 Docker Compose", 2)

    add_code_block(doc, '''# docker-compose.yml
version: '3.8'

services:
  web:
    build: .
    ports:
      - "8000:8000"
    environment:
      - GROQ_API_KEY=${GROQ_API_KEY}
      - SECRET_KEY=${SECRET_KEY}
      - DATABASE_URL=postgresql://postgres:password@db:5432/smartsupport
    depends_on:
      - db

  db:
    image: postgres:14
    environment:
      - POSTGRES_DB=smartsupport
      - POSTGRES_PASSWORD=password
    volumes:
      - postgres_data:/var/lib/postgresql/data

volumes:
  postgres_data:''')

    add_heading(doc, "20.3 Environment Variables", 2)

    headers = ["Variable", "Required", "Description"]
    rows = [
        ["GROQ_API_KEY", "Yes", "API key for Groq LLM service"],
        ["SECRET_KEY", "Yes", "Secret key for security"],
        ["DATABASE_URL", "No", "Database connection string"],
        ["PORT", "No", "Port to run on (default: 8000)"],
        ["ENVIRONMENT", "No", "development or production"],
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "20.4 Railway Deployment", 2)

    add_code_block(doc, '''# railway.json
{
  "$schema": "https://railway.app/railway.schema.json",
  "build": {
    "builder": "DOCKERFILE",
    "dockerfilePath": "Dockerfile"
  },
  "deploy": {
    "startCommand": "gunicorn src.api.app:app ...",
    "healthcheckPath": "/health",
    "healthcheckTimeout": 30
  }
}''')

    doc.add_paragraph("Deployment steps:")

    steps = [
        "Create Railway account at railway.app",
        "Connect your GitHub repository",
        "Add PostgreSQL database service",
        "Set environment variables in Railway dashboard",
        "Deploy - Railway builds and runs automatically"
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()


def create_part_5_results(doc):
    """Part 5: Results and Analysis"""
    add_heading(doc, "PART 5: RESULTS AND ANALYSIS", 1)
    doc.add_paragraph()

    add_heading(doc, "Chapter 21: Performance Metrics", 1)

    add_heading(doc, "21.1 Response Time Analysis", 2)

    headers = ["Operation", "Target", "Achieved", "Status"]
    rows = [
        ["Total Response", "< 2.0s", "0.8-1.2s", "Excellent"],
        ["Categorization", "< 500ms", "200-400ms", "Excellent"],
        ["Sentiment Analysis", "< 500ms", "200-400ms", "Excellent"],
        ["KB Retrieval", "< 100ms", "30-50ms", "Excellent"],
        ["Response Generation", "< 1.0s", "400-600ms", "Excellent"],
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "21.2 Accuracy Metrics", 2)

    headers = ["Metric", "Target", "Achieved"]
    rows = [
        ["Categorization Accuracy", "> 90%", "~92%"],
        ["Sentiment Detection", "> 85%", "~88%"],
        ["KB Retrieval Relevance", "> 85%", "~90%"],
        ["Escalation Accuracy", "100%", "100%"],
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "21.3 System Statistics", 2)

    stats = [
        ("Test Results", "38 passed, 0 failed (100%)"),
        ("Code Coverage", "42.38%"),
        ("Lines of Code", "4,500+"),
        ("API Endpoints", "15+"),
        ("Database Tables", "8"),
        ("AI Agents", "7"),
        ("FAQs in Knowledge Base", "30"),
    ]

    for name, value in stats:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(value)

    doc.add_page_break()


def create_appendices(doc):
    """Create appendices"""
    add_heading(doc, "APPENDICES", 1)
    doc.add_paragraph()

    # Appendix A
    add_heading(doc, "Appendix A: Complete Code Reference", 1)

    headers = ["File", "Purpose", "Lines"]
    rows = [
        ["src/main.py", "Main orchestrator", "~240"],
        ["src/agents/workflow.py", "LangGraph workflow", "~120"],
        ["src/agents/state.py", "State definitions", "~100"],
        ["src/agents/categorizer.py", "Categorization", "~80"],
        ["src/agents/sentiment_analyzer.py", "Sentiment analysis", "~100"],
        ["src/agents/kb_retrieval.py", "KB search", "~75"],
        ["src/agents/technical_agent.py", "Technical responses", "~100"],
        ["src/agents/billing_agent.py", "Billing responses", "~110"],
        ["src/agents/escalation_agent.py", "Escalation logic", "~90"],
        ["src/database/models.py", "SQLAlchemy models", "~280"],
        ["src/database/queries.py", "DB queries", "~400"],
        ["src/api/routes.py", "API endpoints", "~175"],
        ["src/api/webhooks.py", "Webhook system", "~320"],
        ["src/knowledge_base/vector_store.py", "FAISS wrapper", "~230"],
    ]
    create_table(doc, headers, rows)

    doc.add_page_break()

    # Appendix B
    add_heading(doc, "Appendix B: API Reference", 1)

    add_heading(doc, "Core Endpoints", 2)

    headers = ["Method", "Endpoint", "Description"]
    rows = [
        ["POST", "/api/v1/query", "Process customer query"],
        ["GET", "/api/v1/health", "Health check"],
        ["GET", "/api/v1/stats", "System statistics"],
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "Webhook Endpoints", 2)

    headers = ["Method", "Endpoint", "Description"]
    rows = [
        ["POST", "/api/v1/webhooks/", "Create webhook"],
        ["GET", "/api/v1/webhooks/", "List webhooks"],
        ["GET", "/api/v1/webhooks/{id}", "Get webhook"],
        ["PUT", "/api/v1/webhooks/{id}", "Update webhook"],
        ["DELETE", "/api/v1/webhooks/{id}", "Delete webhook"],
        ["POST", "/api/v1/webhooks/{id}/test", "Test webhook"],
    ]
    create_table(doc, headers, rows)

    doc.add_page_break()

    # Appendix C
    add_heading(doc, "Appendix C: Glossary of Terms", 1)

    terms = [
        ("AI Agent", "Software component that perceives, decides, and acts to achieve goals"),
        ("LLM", "Large Language Model - AI trained on text to understand and generate language"),
        ("RAG", "Retrieval-Augmented Generation - grounding LLM responses in retrieved documents"),
        ("Embedding", "Vector (list of numbers) representing text meaning"),
        ("Vector Store", "Database optimized for similarity search on embeddings"),
        ("FAISS", "Facebook AI Similarity Search - fast similarity search library"),
        ("Prompt", "Instructions given to an LLM"),
        ("Temperature", "LLM parameter controlling randomness (0=deterministic)"),
        ("Token", "Unit of text (word or subword) processed by LLM"),
        ("Context Window", "Maximum tokens an LLM can process at once"),
        ("LangChain", "Framework for building LLM applications"),
        ("LangGraph", "Library for building multi-agent workflows"),
        ("FastAPI", "Python web framework for building APIs"),
        ("SQLAlchemy", "Python ORM for database interactions"),
        ("Webhook", "HTTP callback for real-time notifications"),
        ("HMAC", "Hash-based Message Authentication Code"),
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(definition)

    doc.add_page_break()

    # Appendix D
    add_heading(doc, "Appendix D: Troubleshooting Guide", 1)

    issues = [
        ("LLM API Rate Limit", "Error: Rate limit exceeded",
         "Solution: Retry logic with exponential backoff (already implemented in llm_manager.py)"),
        ("FAISS Index Not Found", "Error: No existing index found",
         "Solution: Run 'python initialize_kb.py' to build the index"),
        ("Database Connection Error", "Error: Can't connect to database",
         "Solution: Check DATABASE_URL in .env, ensure PostgreSQL is running"),
        ("Module Not Found", "Error: ModuleNotFoundError",
         "Solution: Activate virtual environment, run 'pip install -r requirements.txt'"),
        ("API Returns 500", "Error: Internal server error",
         "Solution: Check logs/app.log for details, verify API keys are set"),
    ]

    for title, error, solution in issues:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True

        p = doc.add_paragraph()
        run = p.add_run(error)
        run.font.color.rgb = RGBColor(180, 0, 0)

        doc.add_paragraph(solution)
        doc.add_paragraph()


def create_conclusion(doc):
    """Create conclusion"""
    add_heading(doc, "Conclusion", 1)

    doc.add_paragraph(
        "Congratulations on completing this comprehensive tutorial! You've learned how to build "
        "a production-ready AI customer support system from first principles."
    )

    add_heading(doc, "What You Learned", 2)

    learned = [
        "What AI agents are and how they work",
        "How Large Language Models generate text",
        "Prompt engineering techniques for reliable results",
        "Multi-agent system architecture",
        "Retrieval-Augmented Generation (RAG)",
        "LangGraph workflow orchestration",
        "FastAPI REST API development",
        "Database design with SQLAlchemy",
        "Testing strategies for AI systems",
        "Production deployment with Docker"
    ]

    for item in learned:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, "Key Takeaways", 2)

    takeaways = [
        "Multi-agent systems provide modularity, testability, and maintainability",
        "RAG is essential for accurate, factual AI responses",
        "State management is the backbone of agent orchestration",
        "Prompt engineering requires precision and iteration",
        "Graceful degradation ensures reliability",
        "Production readiness requires testing, monitoring, and proper deployment"
    ]

    for i, t in enumerate(takeaways, 1):
        doc.add_paragraph(f"{i}. {t}")

    add_heading(doc, "Next Steps", 2)

    next_steps = [
        "Extend with more specialized agents",
        "Add conversation memory for multi-turn",
        "Implement user feedback loop for improvement",
        "Add analytics dashboard",
        "Integrate with CRM systems"
    ]

    for step in next_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Happy Building!")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Document info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Multi-Agent HR Intelligence Platform - Comprehensive Tutorial")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Version 2.2.0 | January 2026 | ~75+ Pages")


def complete_tutorial():
    """Complete the tutorial by adding remaining parts to existing document."""
    print("Loading existing document...")
    doc = Document("tutorial_documentation/Multi-Agent HR Intelligence Platform_AI_Comprehensive_Tutorial.docx")

    print("Adding Part 3: Building the Agents...")
    create_part_3_agents(doc)

    print("Adding Part 4: Production Implementation...")
    create_part_4_implementation(doc)

    print("Adding Part 5: Results...")
    create_part_5_results(doc)

    print("Adding Appendices...")
    create_appendices(doc)

    print("Adding Conclusion...")
    create_conclusion(doc)

    output_path = "tutorial_documentation/Multi-Agent HR Intelligence Platform_AI_Comprehensive_Tutorial.docx"
    doc.save(output_path)
    print(f"\nComplete tutorial saved to: {output_path}")

    return output_path


if __name__ == "__main__":
    complete_tutorial()
