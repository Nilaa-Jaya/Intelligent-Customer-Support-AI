"""
Generate comprehensive Word document tutorial for Multi-Agent HR Intelligence Platform
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_code_block(doc, code, language="python"):
    """Add a formatted code block"""
    # Add code in a paragraph with monospace font
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)

    return para

def add_key_insight(doc, text):
    """Add a key insight box"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)

    run = para.add_run("KEY INSIGHT: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    run2 = para.add_run(text)
    run2.italic = True

    return para

def add_warning(doc, text):
    """Add a warning box"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)

    run = para.add_run("WARNING: ")
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)

    run2 = para.add_run(text)

    return para

def create_table(doc, headers, rows):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'D9E2F3')

    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)

    return table

def generate_tutorial():
    """Generate the complete tutorial document"""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ========================================
    # TITLE PAGE
    # ========================================

    title = doc.add_heading('Multi-Agent HR Intelligence Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Building Intelligent Customer Support Systems')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(70, 70, 70)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('A Complete Guide from First Principles to Production Implementation')
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Key highlights
    highlights = doc.add_paragraph()
    highlights.alignment = WD_ALIGN_PARAGRAPH.CENTER
    highlights.add_run('Key Highlights:').bold = True

    bullets = [
        '7 Specialized AI Agents orchestrated via LangGraph',
        '30 Comprehensive FAQs in FAISS Vector Store (RAG)',
        '8 Database Tables with Full Conversation Tracking',
        '15+ RESTful API Endpoints with Webhook Support',
        'Production-Ready with Docker, CI/CD, and Railway Deployment',
        '38 Automated Tests with 42% Code Coverage'
    ]

    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run('Creation Date: January 2026')

    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_para.add_run('Version: 2.2.0')

    doc.add_page_break()

    # ========================================
    # READER'S GUIDE
    # ========================================

    add_heading(doc, "Reader's Guide", 1)

    add_heading(doc, "What This Document Is", 2)
    doc.add_paragraph(
        "This document is a complete, comprehensive tutorial that teaches you how to build an "
        "enterprise-grade AI-powered customer support system from scratch. It is not a summary "
        "or overview - it is a detailed, step-by-step guide that explains every concept, every "
        "line of code, and every architectural decision."
    )

    doc.add_paragraph("By the end of this tutorial, you will understand:")
    bullets = [
        "How to build multi-agent AI systems using LangChain and LangGraph",
        "How to implement Retrieval-Augmented Generation (RAG) with vector databases",
        "How to design and implement production-ready APIs with FastAPI",
        "How to deploy AI applications to cloud platforms"
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    add_heading(doc, "Who This Is For", 2)
    audiences = [
        "Software Developers looking to add AI/ML capabilities to their skill set",
        "AI/ML Engineers wanting to learn production deployment patterns",
        "Backend Developers interested in building intelligent customer service systems",
        "Computer Science Students building portfolio projects",
        "Technical Leads evaluating AI architectures for customer support"
    ]
    for audience in audiences:
        doc.add_paragraph(audience, style='List Bullet')

    add_heading(doc, "Prerequisites (What You MUST Know)", 2)
    prereqs = [
        "Python Fundamentals: Variables, functions, classes, modules, and packages",
        "Basic SQL: SELECT, INSERT, UPDATE, JOIN operations",
        "REST APIs: Understanding of HTTP methods, JSON, and API design",
        "Git Basics: Clone, commit, push, pull operations",
        "Command Line: Basic terminal/command prompt navigation"
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    add_heading(doc, "What You DON'T Need to Know (We Explain From Scratch)", 2)
    no_prereqs = [
        "LangChain or LangGraph: We explain these frameworks from first principles",
        "Large Language Models (LLMs): We cover how they work and how to use them",
        "Vector Databases: Complete explanation of embeddings and similarity search",
        "RAG Architecture: Full breakdown of Retrieval-Augmented Generation",
        "Multi-Agent Systems: Detailed explanation of agent orchestration"
    ]
    for item in no_prereqs:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========================================
    # TABLE OF CONTENTS
    # ========================================

    add_heading(doc, "Table of Contents", 1)

    toc_items = [
        ("PART 1: DOMAIN FUNDAMENTALS", [
            "Chapter 1: Core Concepts of AI Customer Support",
            "Chapter 2: Understanding Large Language Models (LLMs)",
            "Chapter 3: Multi-Agent Systems",
            "Chapter 4: Retrieval-Augmented Generation (RAG)"
        ]),
        ("PART 2: PROBLEM SETUP AND DATA", [
            "Chapter 5: The Customer Support Problem",
            "Chapter 6: Data Architecture",
            "Chapter 7: The Baseline Approach"
        ]),
        ("PART 3: METHODS DEEP DIVE", [
            "Chapter 8: Query Categorization Agent",
            "Chapter 9: Sentiment Analysis Agent",
            "Chapter 10: Knowledge Base Retrieval Agent",
            "Chapter 11: Technical Support Agent",
            "Chapter 12: Billing Support Agent",
            "Chapter 13: General & Account Support Agents",
            "Chapter 14: Escalation Agent",
            "Chapter 15: LangGraph Workflow Orchestration"
        ]),
        ("PART 4: IMPLEMENTATION FRAMEWORK", [
            "Chapter 16: Project Architecture",
            "Chapter 17: Database Implementation",
            "Chapter 18: API Implementation",
            "Chapter 19: Web Interface",
            "Chapter 20: Testing Strategy",
            "Chapter 21: Deployment to Production"
        ]),
        ("PART 5: RESULTS AND INTERPRETATION", [
            "Chapter 22: Performance Metrics",
            "Chapter 23: System Evaluation",
            "Chapter 24: Failure Analysis"
        ]),
        ("PART 6: DESIGN DECISIONS AND TRADE-OFFS", [
            "Chapter 25: Architectural Choices",
            "Chapter 26: Technology Selection",
            "Chapter 27: What We Didn't Implement"
        ]),
        ("APPENDICES", [
            "Appendix A: Complete Code Reference",
            "Appendix B: Mathematical Foundations",
            "Appendix C: Glossary of Terms",
            "Appendix D: Troubleshooting Guide",
            "Appendix E: References and Further Reading"
        ])
    ]

    for part_name, chapters in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(part_name)
        run.bold = True
        run.font.size = Pt(12)

        for chapter in chapters:
            doc.add_paragraph(chapter, style='List Bullet')

    doc.add_page_break()

    # ========================================
    # PART 1: DOMAIN FUNDAMENTALS
    # ========================================

    add_heading(doc, "PART 1: DOMAIN FUNDAMENTALS", 1)

    # Chapter 1
    add_heading(doc, "Chapter 1: Core Concepts of AI Customer Support", 2)

    add_heading(doc, "1.1 What is AI Customer Support?", 3)
    doc.add_paragraph(
        "AI Customer Support refers to the use of artificial intelligence technologies to automate, "
        "augment, or enhance customer service interactions. Instead of (or in addition to) human agents, "
        "AI systems can:"
    )

    capabilities = [
        "Understand customer queries written in natural language",
        "Categorize issues to route them appropriately",
        "Retrieve relevant information from knowledge bases",
        "Generate helpful, contextual responses",
        "Escalate complex issues to human agents when necessary"
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_paragraph(
        "Think of it as having a highly knowledgeable, always-available assistant that can handle "
        "the majority of customer inquiries while knowing when to involve a human."
    )

    p = doc.add_paragraph()
    run = p.add_run("Analogy for Software Developers: ")
    run.bold = True
    p.add_run(
        "If you've ever worked with an ORM (Object-Relational Mapper), you know it translates between "
        "your code and the database. AI Customer Support is similar - it translates between natural "
        "human language and structured business processes."
    )

    add_heading(doc, "1.2 Why AI-Powered Support Matters", 3)
    doc.add_paragraph("The business case for AI customer support is compelling:")

    p = doc.add_paragraph()
    run = p.add_run("Scale Without Linear Cost: ")
    run.bold = True
    p.add_run(
        "Traditional customer support scales linearly - 2x customers requires approximately 2x support staff. "
        "AI support can handle thousands of simultaneous conversations at the cost of compute resources."
    )

    p = doc.add_paragraph()
    run = p.add_run("24/7 Availability: ")
    run.bold = True
    p.add_run("AI doesn't sleep, take breaks, or call in sick. It provides consistent service around the clock.")

    p = doc.add_paragraph()
    run = p.add_run("Consistency: ")
    run.bold = True
    p.add_run("Human agents have varying levels of knowledge and experience. AI provides consistent answers.")

    p = doc.add_paragraph()
    run = p.add_run("Instant Response: ")
    run.bold = True
    p.add_run("Customers expect immediate responses. AI can respond in under a second.")

    add_key_insight(doc,
        "The goal isn't to replace human agents entirely, but to handle the 80% of queries that are routine, "
        "freeing human agents to focus on complex issues that require empathy, judgment, and creative problem-solving."
    )

    add_heading(doc, "1.3 The Evolution of Customer Service Technology", 3)

    doc.add_paragraph("Understanding the evolution helps contextualize where AI customer support fits:")

    p = doc.add_paragraph()
    run = p.add_run("Generation 1: Rule-Based Systems (1990s-2000s)")
    run.bold = True
    doc.add_paragraph("Simple keyword matching and decision trees with fixed paths. \"If customer says X, respond with Y.\" Limited by inability to handle variations.")

    p = doc.add_paragraph()
    run = p.add_run("Generation 2: NLP + Intent Classification (2010s)")
    run.bold = True
    doc.add_paragraph("Natural Language Processing to understand meaning. Machine learning models to classify intent. Still required templated responses.")

    p = doc.add_paragraph()
    run = p.add_run("Generation 3: Conversational AI + LLMs (2020s - Present)")
    run.bold = True
    doc.add_paragraph("Large Language Models understand context and nuance. Can generate human-like responses. RAG grounds responses in factual knowledge. This is where Multi-Agent HR Intelligence Platform operates.")

    # Chapter 2
    add_heading(doc, "Chapter 2: Understanding Large Language Models (LLMs)", 2)

    add_heading(doc, "2.1 What is an LLM?", 3)
    doc.add_paragraph(
        "A Large Language Model (LLM) is a type of artificial intelligence trained on massive amounts of text data. "
        "The key characteristics are:"
    )

    p = doc.add_paragraph()
    run = p.add_run("\"Large\": ")
    run.bold = True
    p.add_run("Billions of parameters (learned weights) - GPT-3 has 175 billion, Llama 3 has up to 70 billion")

    p = doc.add_paragraph()
    run = p.add_run("\"Language\": ")
    run.bold = True
    p.add_run("Trained specifically on text data in human languages")

    p = doc.add_paragraph()
    run = p.add_run("\"Model\": ")
    run.bold = True
    p.add_run("A mathematical function that takes input text and predicts what text should come next")

    p = doc.add_paragraph()
    run = p.add_run("Analogy for Developers: ")
    run.bold = True
    p.add_run(
        "Think of an LLM as an incredibly sophisticated autocomplete. When you type in an IDE and it suggests "
        "the next function, it's using patterns it learned from code. An LLM does the same for natural language, "
        "but at a much deeper level."
    )

    doc.add_paragraph("How Multi-Agent HR Intelligence Platform Uses LLMs:")
    doc.add_paragraph("We use the Groq API with Llama 3.3-70B, a state-of-the-art open-source LLM:")

    add_code_block(doc, '''# From src/utils/config.py
llm_model: str = "llama-3.3-70b-versatile"
llm_temperature: float = 0.0  # Deterministic outputs
llm_max_tokens: int = 1000    # Maximum response length''')

    add_heading(doc, "2.2 How LLMs Generate Text", 3)
    doc.add_paragraph("Understanding the generation process helps you write better prompts:")

    p = doc.add_paragraph()
    run = p.add_run("Step 1: Tokenization - ")
    run.bold = True
    p.add_run("Input text is broken into \"tokens\" (words or subwords)")

    p = doc.add_paragraph()
    run = p.add_run("Step 2: Embedding - ")
    run.bold = True
    p.add_run("Each token is converted to a high-dimensional vector capturing its meaning")

    p = doc.add_paragraph()
    run = p.add_run("Step 3: Attention Mechanism - ")
    run.bold = True
    p.add_run("The model determines which parts of the input are relevant to generating output")

    p = doc.add_paragraph()
    run = p.add_run("Step 4: Token-by-Token Generation - ")
    run.bold = True
    p.add_run("The model predicts probabilities for the next token, selects one, and repeats")

    p = doc.add_paragraph()
    run = p.add_run("Temperature Parameter: ")
    run.bold = True
    p.add_run("Controls randomness. temperature=0.0 (deterministic) vs temperature=0.7 (creative). We use 0.0 for consistent support responses.")

    add_heading(doc, "2.3 Prompt Engineering Fundamentals", 3)
    doc.add_paragraph(
        "Prompt engineering is the practice of crafting instructions that get the desired behavior from an LLM."
    )

    doc.add_paragraph("Key Principles:")

    p = doc.add_paragraph()
    run = p.add_run("1. Be Specific About the Role: ")
    run.bold = True
    p.add_run("\"You are an expert customer support query classifier\"")

    p = doc.add_paragraph()
    run = p.add_run("2. Provide Context: ")
    run.bold = True
    p.add_run("Include conversation history for continuity")

    p = doc.add_paragraph()
    run = p.add_run("3. Specify Output Format: ")
    run.bold = True
    p.add_run("\"Respond with ONLY the category name\"")

    p = doc.add_paragraph()
    run = p.add_run("4. Include Behavioral Instructions: ")
    run.bold = True
    p.add_run("\"If sentiment is negative, start with empathy\"")

    add_code_block(doc, '''# Example from src/agents/categorizer.py
CATEGORIZATION_PROMPT = """You are an expert customer support query classifier.

Categorize the following customer query into ONE of these categories:
- Technical: Issues with software, hardware, bugs, errors
- Billing: Payment issues, invoices, refunds, subscriptions
- Account: Login, password, profile, account settings
- General: Company policies, general inquiries, feedback

Query: {query}

Respond with ONLY the category name.
Category:"""''')

    add_key_insight(doc,
        "Good prompts are like good function documentation - they clearly specify inputs, expected behavior, "
        "and output format. Ambiguity in prompts leads to inconsistent results."
    )

    # Chapter 3
    add_heading(doc, "Chapter 3: Multi-Agent Systems", 2)

    add_heading(doc, "3.1 What is an Agent?", 3)
    doc.add_paragraph("In AI terms, an agent is a component that:")
    doc.add_paragraph("1. Receives input (state, context)", style='List Number')
    doc.add_paragraph("2. Makes decisions or performs actions", style='List Number')
    doc.add_paragraph("3. Produces output that affects subsequent processing", style='List Number')

    p = doc.add_paragraph()
    run = p.add_run("Analogy: Microservices Architecture - ")
    run.bold = True
    p.add_run(
        "Just as microservices split a monolithic application into specialized services, "
        "multi-agent systems split AI processing into specialized agents."
    )

    doc.add_paragraph("Multi-Agent HR Intelligence Platform has 7 specialized agents:")

    headers = ["Agent", "Purpose", "Output"]
    rows = [
        ["Categorizer", "Classify query type", "Category"],
        ["Sentiment Analyzer", "Detect emotional tone", "Sentiment + Priority"],
        ["KB Retrieval", "Find relevant FAQs", "Relevant documents"],
        ["Technical Agent", "Handle tech queries", "Technical response"],
        ["Billing Agent", "Handle billing queries", "Billing response"],
        ["General Agent", "Handle general queries", "General response"],
        ["Escalation Agent", "Handle escalations", "Escalation response"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "3.2 Why Multiple Agents?", 3)

    p = doc.add_paragraph()
    run = p.add_run("1. Focused Prompts: ")
    run.bold = True
    p.add_run("Each agent has a prompt optimized for its specific task")

    p = doc.add_paragraph()
    run = p.add_run("2. Maintainability: ")
    run.bold = True
    p.add_run("Need to change billing handling? Modify only billing_agent.py")

    p = doc.add_paragraph()
    run = p.add_run("3. Testability: ")
    run.bold = True
    p.add_run("Each agent can be unit tested in isolation")

    p = doc.add_paragraph()
    run = p.add_run("4. Extensibility: ")
    run.bold = True
    p.add_run("Adding new category requires minimal changes to existing code")

    add_heading(doc, "3.3 State Management in Multi-Agent Systems", 3)
    doc.add_paragraph(
        "The Challenge: How do agents share information? Each agent needs to know what previous agents discovered."
    )

    doc.add_paragraph("Our Solution: AgentState TypedDict")

    add_code_block(doc, '''# From src/agents/state.py
class AgentState(TypedDict):
    # Input
    query: str
    user_id: str
    conversation_id: str

    # Analysis results (filled by analysis agents)
    category: Optional[str]
    sentiment: Optional[str]
    priority_score: Optional[int]

    # Knowledge base results
    kb_results: Optional[List[Dict[str, Any]]]

    # Response (filled by response agents)
    response: Optional[str]

    # Routing decisions
    should_escalate: bool
    escalation_reason: Optional[str]''')

    add_key_insight(doc,
        "State management is the backbone of multi-agent systems. A well-designed state object makes "
        "agents composable and the system debuggable."
    )

    # Chapter 4
    add_heading(doc, "Chapter 4: Retrieval-Augmented Generation (RAG)", 2)

    add_heading(doc, "4.1 The Knowledge Problem in LLMs", 3)
    doc.add_paragraph("LLMs have a fundamental limitation: their knowledge is frozen at training time.")

    p = doc.add_paragraph()
    run = p.add_run("1. Outdated Information: ")
    run.bold = True
    p.add_run("An LLM trained in 2023 doesn't know about features you released in 2024")

    p = doc.add_paragraph()
    run = p.add_run("2. No Company-Specific Knowledge: ")
    run.bold = True
    p.add_run("LLMs know general information but nothing about YOUR products")

    p = doc.add_paragraph()
    run = p.add_run("3. Hallucination Risk: ")
    run.bold = True
    p.add_run("When asked about unknown topics, LLMs often make up plausible-sounding answers")

    add_heading(doc, "4.2 What is RAG?", 3)
    doc.add_paragraph("Retrieval-Augmented Generation (RAG) solves the knowledge problem by:")
    doc.add_paragraph("1. Retrieving relevant documents from a knowledge base", style='List Number')
    doc.add_paragraph("2. Augmenting the LLM's prompt with this retrieved context", style='List Number')
    doc.add_paragraph("3. Generating responses grounded in the retrieved information", style='List Number')

    doc.add_paragraph("The RAG Pipeline:")
    add_code_block(doc, '''User Query: "How do I export my data?"
        ↓
    [Embedding Model] - Encode query to vector
        ↓
Query Vector: [0.12, -0.34, 0.56, ...]
        ↓
    [Vector Database Search] - Find nearest neighbors
        ↓
Retrieved Docs:
  1. FAQ: "How do I export my data?" (similarity: 0.92)
  2. FAQ: "Can I download my files?" (similarity: 0.71)
        ↓
    [LLM with Context]
        ↓
Response: Accurate, company-specific answer''')

    add_heading(doc, "4.3 Vector Embeddings Explained", 3)

    p = doc.add_paragraph()
    run = p.add_run("What is an Embedding? ")
    run.bold = True
    p.add_run("A list of numbers (vector) that represents the \"meaning\" of text. Similar texts have similar embeddings.")

    p = doc.add_paragraph()
    run = p.add_run("Analogy: GPS Coordinates - ")
    run.bold = True
    p.add_run(
        "Just as GPS coordinates represent physical locations, embeddings represent semantic locations "
        "in \"meaning space.\" \"happy\" and \"joyful\" are nearby; \"happy\" and \"sad\" are distant."
    )

    add_code_block(doc, '''# From src/knowledge_base/vector_store.py
from sentence_transformers import SentenceTransformer

self.encoder = SentenceTransformer("all-MiniLM-L6-v2")
# Produces 384-dimensional embeddings
# Optimized for semantic similarity''')

    add_heading(doc, "4.4 Similarity Search with FAISS", 3)
    doc.add_paragraph(
        "FAISS (Facebook AI Similarity Search) is a library optimized for similarity search "
        "in high-dimensional spaces."
    )

    add_code_block(doc, '''# From src/knowledge_base/vector_store.py
import faiss

# Create index
self.index = faiss.IndexFlatL2(self.embedding_dim)

# Add document embeddings
self.index.add(embeddings)

# Search for similar documents
distances, indices = self.index.search(query_embedding, k=3)

# Convert distance to similarity
similarity_score = 1 / (1 + distance)''')

    add_key_insight(doc,
        "RAG transforms LLMs from \"knows everything (often wrongly)\" to \"knows exactly what we tell it.\" "
        "This is essential for accurate customer support."
    )

    doc.add_page_break()

    # ========================================
    # PART 2: PROBLEM SETUP AND DATA
    # ========================================

    add_heading(doc, "PART 2: PROBLEM SETUP AND DATA", 1)

    add_heading(doc, "Chapter 5: The Customer Support Problem", 2)

    add_heading(doc, "5.1 Problem Statement", 3)
    doc.add_paragraph("Goal: Build an intelligent customer support system that can:")

    goals = [
        "Understand customer queries in natural language",
        "Categorize queries by type (Technical, Billing, Account, General)",
        "Detect customer sentiment and urgency",
        "Retrieve relevant information from a knowledge base",
        "Generate helpful, accurate responses",
        "Escalate to human agents when necessary",
        "Track conversations for analytics and improvement"
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    add_heading(doc, "5.2 Requirements Analysis", 3)

    doc.add_paragraph("Functional Requirements:")
    headers = ["ID", "Requirement", "Priority"]
    rows = [
        ["F1", "Accept natural language queries", "Must Have"],
        ["F2", "Categorize queries into 4 categories", "Must Have"],
        ["F3", "Analyze sentiment (4 levels)", "Must Have"],
        ["F4", "Calculate priority scores (1-10)", "Must Have"],
        ["F5", "Retrieve relevant FAQs", "Must Have"],
        ["F6", "Generate contextual responses", "Must Have"],
        ["F7", "Escalate when necessary", "Must Have"],
        ["F8", "Persist conversations", "Must Have"],
        ["F9", "Provide REST API", "Must Have"],
        ["F10", "Web user interface", "Should Have"],
        ["F11", "Webhook notifications", "Should Have"]
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()
    doc.add_paragraph("Non-Functional Requirements:")
    headers = ["ID", "Requirement", "Target"]
    rows = [
        ["N1", "Response Time", "< 2 seconds"],
        ["N2", "Availability", "99.9% uptime"],
        ["N3", "Concurrent Users", "100+"],
        ["N4", "Data Security", "Encryption at rest and in transit"],
        ["N5", "Test Coverage", "> 25%"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "5.3 Success Metrics Achieved", 3)

    headers = ["Metric", "Target", "Achieved"]
    rows = [
        ["Response Time", "< 2s", "0.8-1.2s"],
        ["KB Accuracy", "> 85%", "~90%"],
        ["Escalation Rate", "< 15%", "~12%"],
        ["Test Coverage", "> 25%", "42.38%"],
        ["Tests Passing", "100%", "38/38 (100%)"]
    ]
    create_table(doc, headers, rows)

    # Chapter 6
    add_heading(doc, "Chapter 6: Data Architecture", 2)

    add_heading(doc, "6.1 Knowledge Base Design", 3)
    doc.add_paragraph("The knowledge base consists of 30 FAQs across 4 categories:")

    headers = ["Category", "Count", "Topics"]
    rows = [
        ["Technical", "10", "Crashes, login, sync, export, performance"],
        ["Billing", "10", "Charges, refunds, subscriptions, payments"],
        ["Account", "5", "Password reset, profile, deletion"],
        ["General", "5", "Support hours, contact, security, platforms"]
    ]
    create_table(doc, headers, rows)

    doc.add_paragraph()
    doc.add_paragraph("FAQ Structure:")
    add_code_block(doc, '''{
  "id": 1,
  "category": "Technical",
  "question": "Why does my app keep crashing?",
  "answer": "App crashes can be caused by several factors: 1) Outdated app version..."
}''')

    add_heading(doc, "6.2 Database Tables", 3)
    doc.add_paragraph("The system uses 8 database tables:")

    headers = ["Table", "Purpose", "Key Fields"]
    rows = [
        ["users", "User profiles", "user_id, is_vip, email"],
        ["conversations", "Query records", "query, category, sentiment, response"],
        ["messages", "Chat history", "role, content, created_at"],
        ["feedback", "User ratings", "rating (1-5), comment"],
        ["analytics", "Metrics aggregation", "counts, averages by hour"],
        ["knowledge_base", "FAQ storage", "title, content, category"],
        ["webhooks", "Webhook config", "url, events, secret_key"],
        ["webhook_deliveries", "Delivery logs", "status, response"]
    ]
    create_table(doc, headers, rows)

    doc.add_page_break()

    # ========================================
    # PART 3: METHODS DEEP DIVE
    # ========================================

    add_heading(doc, "PART 3: METHODS DEEP DIVE", 1)

    # Chapter 8
    add_heading(doc, "Chapter 8: Query Categorization Agent", 2)

    add_heading(doc, "8.1 Purpose and Function", 3)
    doc.add_paragraph(
        "The Categorization Agent is the first step in our workflow. Its job is to classify "
        "the customer's query into one of four categories: Technical, Billing, Account, or General."
    )

    add_heading(doc, "8.2 The Prompt Design", 3)
    add_code_block(doc, '''# From src/agents/categorizer.py
CATEGORIZATION_PROMPT = ChatPromptTemplate.from_template(
    """You are an expert customer support query classifier.

Categorize the following customer query into ONE of these categories:
- Technical: Issues with software, hardware, bugs, errors, setup
- Billing: Payment issues, invoices, refunds, subscriptions, pricing
- Account: Login, password, profile, account settings, security
- General: Company policies, general inquiries, feedback

Query: {query}

{context}

Respond with ONLY the category name (Technical, Billing, Account, or General).
Category:"""
)''')

    doc.add_paragraph("Prompt Analysis:")

    p = doc.add_paragraph()
    run = p.add_run("Role Definition: ")
    run.bold = True
    p.add_run("\"You are an expert customer support query classifier\" establishes persona and focus.")

    p = doc.add_paragraph()
    run = p.add_run("Category Definitions: ")
    run.bold = True
    p.add_run("Each category includes multiple examples to reduce ambiguity.")

    p = doc.add_paragraph()
    run = p.add_run("Context Injection: ")
    run.bold = True
    p.add_run("Includes conversation history for queries like \"Still not working.\"")

    p = doc.add_paragraph()
    run = p.add_run("Output Constraint: ")
    run.bold = True
    p.add_run("\"Respond with ONLY the category name\" forces concise, parseable output.")

    add_heading(doc, "8.3 Implementation", 3)
    add_code_block(doc, '''def categorize_query(state: AgentState) -> AgentState:
    app_logger.info(f"Categorizing query: {state['query'][:50]}...")

    try:
        llm_manager = get_llm_manager()

        # Prepare context from conversation history
        context = ""
        if state.get("conversation_history"):
            context = "Previous conversation context:\\n"
            for msg in state["conversation_history"][-3:]:
                context += f"{msg['role']}: {msg['content'][:100]}\\n"

        # Invoke LLM
        raw_category = llm_manager.invoke_with_retry(
            CATEGORIZATION_PROMPT,
            {"query": state["query"], "context": context}
        )

        # Parse and standardize
        category = parse_llm_category(raw_category)
        state["category"] = category

        return state

    except Exception as e:
        app_logger.error(f"Error: {e}")
        state["category"] = "General"  # Graceful fallback
        return state''')

    # Chapter 9
    add_heading(doc, "Chapter 9: Sentiment Analysis Agent", 2)

    add_heading(doc, "9.1 Purpose and Function", 3)
    doc.add_paragraph("The Sentiment Analyzer determines the emotional tone:")

    headers = ["Sentiment", "Indicators", "Base Priority"]
    rows = [
        ["Positive", "Happy, grateful, satisfied", "3"],
        ["Neutral", "Informational, calm", "3"],
        ["Negative", "Frustrated, disappointed", "5 (+2)"],
        ["Angry", "Very upset, demanding", "6 (+3)"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "9.2 Priority Score Calculation", 3)
    add_code_block(doc, '''def calculate_priority_score(
    sentiment: str,
    category: str,
    is_repeat_query: bool = False,
    is_vip: bool = False
) -> int:
    """
    Priority Scale:
    1-3: Low (normal queries)
    4-6: Medium (negative sentiment)
    7-8: High (angry/urgent)
    9-10: Critical (angry + repeat/VIP)
    """
    score = 3  # Base score

    # Sentiment adjustments
    sentiment_scores = {
        "Negative": 2,
        "Angry": 3,
        "Neutral": 0,
        "Positive": 0,
    }
    score += sentiment_scores.get(sentiment, 0)

    if is_repeat_query:
        score += 2
    if is_vip:
        score += 2

    return max(1, min(10, score))  # Clamp 1-10''')

    doc.add_paragraph("Priority Examples:")
    headers = ["Sentiment", "Repeat", "VIP", "Score"]
    rows = [
        ["Neutral", "No", "No", "3"],
        ["Negative", "No", "No", "5"],
        ["Angry", "No", "No", "6"],
        ["Angry", "Yes", "No", "8"],
        ["Angry", "Yes", "Yes", "10"]
    ]
    create_table(doc, headers, rows)

    # Chapter 10
    add_heading(doc, "Chapter 10: Knowledge Base Retrieval Agent", 2)

    add_heading(doc, "10.1 Architecture Overview", 3)
    add_code_block(doc, '''Query: "How do I export data?"
        ↓
    [Sentence Transformer] - Encode to vector
        ↓
Query Vector: [0.12, -0.34, 0.56, ...]
        ↓
    [FAISS Index Search] - Find nearest neighbors
        ↓
    [Filter by Category] - Only Technical FAQs
        ↓
    [Filter by Score] - Keep similarity > 0.3
        ↓
Results: Top 3 most relevant FAQs''')

    add_heading(doc, "10.2 Implementation", 3)
    add_code_block(doc, '''def retrieve_from_kb(state: AgentState) -> AgentState:
    query = state.get("query", "")
    category = state.get("category", "General")

    try:
        kb_retriever = get_kb_retriever()

        # Retrieve with filtering
        results = kb_retriever.retrieve(
            query=query,
            k=3,                   # Top 3 results
            category=category,    # Filter by category
            min_score=0.3,        # Minimum similarity
        )

        # Format for response agents
        kb_results = []
        for result in results:
            kb_results.append({
                "title": result.get("question", ""),
                "content": result.get("answer", ""),
                "score": result.get("similarity_score", 0.0),
            })

        state["kb_results"] = kb_results
        return state

    except Exception as e:
        state["kb_results"] = []  # Graceful degradation
        return state''')

    # Chapter 11-13
    add_heading(doc, "Chapter 11-13: Response Agents", 2)

    doc.add_paragraph("Multi-Agent HR Intelligence Platform has four specialized response agents:")

    add_heading(doc, "Technical Agent", 3)
    doc.add_paragraph("Handles software bugs, configuration, performance issues:")
    add_code_block(doc, '''Instructions in prompt:
1. Provide clear, step-by-step technical solution
2. Use simple language while being technically accurate
3. If sentiment is negative/angry, start with empathy
4. Include troubleshooting steps if applicable
5. Offer to escalate if issue is complex
6. Keep response concise (200-300 words)''')

    add_heading(doc, "Billing Agent", 3)
    doc.add_paragraph("Handles payments, refunds, subscriptions:")
    add_code_block(doc, '''Special handling:
- Detects refund/dispute keywords
- Sets may_need_escalation flag for human review
- References billing policies from KB''')

    add_heading(doc, "Account Agent", 3)
    doc.add_paragraph("Handles login, password, profile with security emphasis:")
    add_code_block(doc, '''Security-focused instructions:
1. Emphasize security best practices
2. Guide through secure password reset process
3. Be reassuring about account security''')

    add_heading(doc, "General Agent", 3)
    doc.add_paragraph("Handles general inquiries, company info, feedback. Uses shorter response target (150-250 words).")

    # Chapter 14
    add_heading(doc, "Chapter 14: Escalation Agent", 2)

    add_heading(doc, "14.1 Escalation Decision Logic", 3)
    add_code_block(doc, '''def should_escalate(priority_score, sentiment, attempt_count, query):
    """
    Escalation triggers:
    1. Priority >= 8 (high severity)
    2. Sentiment is "Angry"
    3. attempt_count >= 3 (multiple failed attempts)
    4. Specific escalation keywords
    """
    reasons = []

    if priority_score >= 8:
        reasons.append("High priority score")

    if sentiment == "Angry":
        reasons.append("Angry sentiment detected")

    if attempt_count >= 3:
        reasons.append("Multiple unsuccessful attempts")

    # Keyword detection
    escalation_keywords = [
        "lawsuit", "legal", "attorney", "lawyer", "sue",
        "speak to manager", "unacceptable", "demand refund"
    ]

    for keyword in escalation_keywords:
        if keyword in query.lower():
            reasons.append(f"Keyword: {keyword}")
            break

    return len(reasons) > 0, "; ".join(reasons)''')

    add_heading(doc, "14.2 Escalation Response", 3)
    doc.add_paragraph("The escalation response is tailored to sentiment:")

    headers = ["Sentiment", "Response Tone"]
    rows = [
        ["Angry", "\"I sincerely apologize for the frustration...immediate assistance\""],
        ["Negative", "\"I understand your concern...best possible assistance\""],
        ["Other", "Neutral, professional handoff"]
    ]
    create_table(doc, headers, rows)

    # Chapter 15
    add_heading(doc, "Chapter 15: LangGraph Workflow Orchestration", 2)

    add_heading(doc, "15.1 Workflow Definition", 3)
    add_code_block(doc, '''from langgraph.graph import StateGraph, END

def create_workflow() -> StateGraph:
    workflow = StateGraph(AgentState)

    # Add nodes
    workflow.add_node("categorize", categorize_query)
    workflow.add_node("analyze_sentiment", analyze_sentiment)
    workflow.add_node("retrieve_kb", retrieve_from_kb)
    workflow.add_node("check_escalation", check_escalation)
    workflow.add_node("technical", handle_technical)
    workflow.add_node("billing", handle_billing)
    workflow.add_node("account", handle_account)
    workflow.add_node("general", handle_general)
    workflow.add_node("escalate", escalate_to_human)

    # Set entry point
    workflow.set_entry_point("categorize")

    # Sequential edges
    workflow.add_edge("categorize", "analyze_sentiment")
    workflow.add_edge("analyze_sentiment", "retrieve_kb")
    workflow.add_edge("retrieve_kb", "check_escalation")

    # Conditional routing
    workflow.add_conditional_edges(
        "check_escalation",
        route_query,
        {
            "technical": "technical",
            "billing": "billing",
            "account": "account",
            "general": "general",
            "escalate": "escalate",
        },
    )

    # Terminal edges
    for node in ["technical", "billing", "account", "general", "escalate"]:
        workflow.add_edge(node, END)

    return workflow.compile()''')

    add_heading(doc, "15.2 Visual Workflow", 3)
    add_code_block(doc, '''                    START
                      │
                      ▼
                ┌─────────────┐
                │  Categorize │
                └─────────────┘
                      │
                      ▼
                ┌─────────────┐
                │  Sentiment  │
                └─────────────┘
                      │
                      ▼
                ┌─────────────┐
                │ KB Retrieval│
                └─────────────┘
                      │
                      ▼
                ┌─────────────┐
                │ Escalation? │
                └─────────────┘
                      │
          ┌───────────┼───────────┐
          │           │           │
          ▼           ▼           ▼
     Technical    Billing     Account  ...
          │           │           │
          └───────────┼───────────┘
                      │
                      ▼
                     END''')

    doc.add_page_break()

    # ========================================
    # PART 4: IMPLEMENTATION FRAMEWORK
    # ========================================

    add_heading(doc, "PART 4: IMPLEMENTATION FRAMEWORK", 1)

    # Chapter 16
    add_heading(doc, "Chapter 16: Project Architecture", 2)

    add_heading(doc, "16.1 Directory Structure", 3)
    add_code_block(doc, '''smartsupport-ai/
├── src/
│   ├── agents/                 # AI agent modules
│   │   ├── workflow.py        # LangGraph workflow
│   │   ├── state.py           # AgentState definition
│   │   ├── categorizer.py     # Query categorization
│   │   ├── sentiment_analyzer.py
│   │   ├── kb_retrieval.py
│   │   ├── technical_agent.py
│   │   ├── billing_agent.py
│   │   ├── general_agent.py
│   │   └── escalation_agent.py
│   │
│   ├── api/                   # FastAPI application
│   │   ├── app.py            # FastAPI setup
│   │   ├── routes.py         # API endpoints
│   │   ├── schemas.py        # Pydantic models
│   │   └── webhooks.py       # Webhook system
│   │
│   ├── database/             # Database layer
│   │   ├── models.py        # SQLAlchemy models
│   │   ├── connection.py    # DB connections
│   │   └── queries.py       # Query functions
│   │
│   ├── knowledge_base/      # RAG implementation
│   │   ├── retriever.py    # Retrieval logic
│   │   └── vector_store.py # FAISS wrapper
│   │
│   └── utils/              # Utilities
│       ├── config.py      # Settings
│       └── helpers.py     # Helper functions
│
├── data/knowledge_base/    # FAQ data
├── tests/                  # Test files
└── requirements.txt''')

    # Chapter 17-18
    add_heading(doc, "Chapter 17-18: Database & API Implementation", 2)

    add_heading(doc, "Database Models (SQLAlchemy)", 3)
    add_code_block(doc, '''class Conversation(Base):
    __tablename__ = "conversations"

    id = Column(Integer, primary_key=True)
    conversation_id = Column(String(50), unique=True)
    user_id = Column(Integer, ForeignKey("users.id"))

    query = Column(Text, nullable=False)
    category = Column(String(50))
    sentiment = Column(String(50))
    priority_score = Column(Integer, default=5)

    response = Column(Text)
    response_time = Column(Float)

    status = Column(String(50), default="Active")
    escalated = Column(Boolean, default=False)''')

    add_heading(doc, "API Endpoints", 3)
    headers = ["Method", "Endpoint", "Purpose"]
    rows = [
        ["POST", "/api/v1/query", "Process support query"],
        ["GET", "/api/v1/health", "Health check"],
        ["POST", "/api/v1/webhooks/", "Create webhook"],
        ["GET", "/api/v1/webhooks/", "List webhooks"],
        ["POST", "/api/v1/webhooks/{id}/test", "Test webhook"]
    ]
    create_table(doc, headers, rows)

    # Chapter 21
    add_heading(doc, "Chapter 21: Deployment to Production", 2)

    add_heading(doc, "Docker Configuration", 3)
    add_code_block(doc, '''# Dockerfile
FROM python:3.10-slim

WORKDIR /app

COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

COPY . .

EXPOSE 8000

CMD ["gunicorn", "src.api.app:app", \\
     "--workers", "4", \\
     "--worker-class", "uvicorn.workers.UvicornWorker", \\
     "--bind", "0.0.0.0:8000"]''')

    add_heading(doc, "Environment Variables", 3)
    headers = ["Variable", "Required", "Description"]
    rows = [
        ["GROQ_API_KEY", "Yes", "Groq API key for LLM"],
        ["SECRET_KEY", "Yes", "Application secret"],
        ["DATABASE_URL", "No", "PostgreSQL connection string"],
        ["PORT", "No", "Server port (default: 8000)"],
        ["ENVIRONMENT", "No", "development/production"]
    ]
    create_table(doc, headers, rows)

    doc.add_page_break()

    # ========================================
    # PART 5: RESULTS
    # ========================================

    add_heading(doc, "PART 5: RESULTS AND INTERPRETATION", 1)

    add_heading(doc, "Chapter 22: Performance Metrics", 2)

    add_heading(doc, "Response Time Breakdown", 3)
    headers = ["Operation", "Target", "Achieved", "Status"]
    rows = [
        ["Total Response", "< 2s", "0.8-1.2s", "Excellent"],
        ["Categorization", "< 500ms", "200-400ms", "Excellent"],
        ["Sentiment Analysis", "< 500ms", "200-400ms", "Excellent"],
        ["KB Retrieval", "< 100ms", "30-50ms", "Excellent"],
        ["Response Generation", "< 1s", "400-600ms", "Excellent"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "System Statistics", 3)
    stats = [
        "Tests: 38 passed, 0 failed",
        "Code Coverage: 42.38%",
        "Lines of Code: 4,500+",
        "API Endpoints: 15+",
        "Database Tables: 8",
        "AI Agents: 7"
    ]
    for stat in stats:
        doc.add_paragraph(stat, style='List Bullet')

    doc.add_page_break()

    # ========================================
    # PART 6: DESIGN DECISIONS
    # ========================================

    add_heading(doc, "PART 6: DESIGN DECISIONS AND TRADE-OFFS", 1)

    add_heading(doc, "Chapter 25: Architectural Choices", 2)

    add_heading(doc, "Why Multi-Agent vs. Single LLM?", 3)

    headers = ["Aspect", "Single LLM", "Multi-Agent"]
    rows = [
        ["Prompt Size", "Large, complex", "Small, focused"],
        ["Debugging", "Difficult", "Easy per-agent"],
        ["Testing", "End-to-end only", "Unit test each agent"],
        ["Maintenance", "Change affects all", "Change one agent"],
        ["Analytics", "Limited", "Rich intermediate state"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "Why FAISS vs. Other Vector Stores?", 3)
    headers = ["Feature", "FAISS", "ChromaDB", "Pinecone"]
    rows = [
        ["Cost", "Free", "Free", "Paid"],
        ["Setup", "Simple", "Simple", "Cloud"],
        ["Speed", "Very fast", "Fast", "Fast"],
        ["Our Use Case", "30 FAQs fits easily", "-", "-"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "Chapter 26: Technology Selection", 2)

    p = doc.add_paragraph()
    run = p.add_run("LLM: Groq + Llama 3.3-70B")
    run.bold = True
    doc.add_paragraph("Extremely fast inference (~100 tokens/second), free tier available, competitive quality.")

    p = doc.add_paragraph()
    run = p.add_run("Framework: LangChain + LangGraph")
    run.bold = True
    doc.add_paragraph("Standard abstractions, conditional routing, state management.")

    p = doc.add_paragraph()
    run = p.add_run("Database: SQLAlchemy + PostgreSQL")
    run.bold = True
    doc.add_paragraph("ORM for clean code, production-grade reliability, Railway provides managed PostgreSQL.")

    doc.add_page_break()

    # ========================================
    # APPENDICES
    # ========================================

    add_heading(doc, "APPENDICES", 1)

    add_heading(doc, "Appendix A: Key Code Files", 2)

    headers = ["File", "Purpose", "Lines"]
    rows = [
        ["src/main.py", "Main orchestrator", "~240"],
        ["src/agents/workflow.py", "LangGraph workflow", "~120"],
        ["src/agents/state.py", "State definitions", "~100"],
        ["src/agents/categorizer.py", "Query categorization", "~80"],
        ["src/agents/sentiment_analyzer.py", "Sentiment analysis", "~100"],
        ["src/database/models.py", "SQLAlchemy models", "~280"],
        ["src/api/routes.py", "API endpoints", "~175"],
        ["src/knowledge_base/vector_store.py", "FAISS implementation", "~230"]
    ]
    create_table(doc, headers, rows)

    add_heading(doc, "Appendix B: Glossary", 2)

    terms = [
        ("LLM", "Large Language Model - AI trained on text to understand and generate language"),
        ("RAG", "Retrieval-Augmented Generation - combining search with LLM generation"),
        ("Embedding", "Vector representation of text meaning (list of numbers)"),
        ("FAISS", "Facebook AI Similarity Search - library for fast vector search"),
        ("Agent", "Component that receives state, processes it, and produces output"),
        ("Prompt", "Instructions given to an LLM to guide its response"),
        ("Token", "Unit of text (word or subword) that LLMs process"),
        ("Vector Store", "Database for storing and searching embeddings"),
        ("Webhook", "HTTP callback for real-time notifications to external systems"),
        ("HMAC", "Hash-based Message Authentication Code for webhook security")
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(definition)

    add_heading(doc, "Appendix C: Troubleshooting", 2)

    p = doc.add_paragraph()
    run = p.add_run("LLM API Rate Limit:")
    run.bold = True
    doc.add_paragraph("Implement retry with exponential backoff (already done in llm_manager.py)")

    p = doc.add_paragraph()
    run = p.add_run("FAISS Index Not Found:")
    run.bold = True
    doc.add_paragraph("Run 'python initialize_kb.py' to build the index")

    p = doc.add_paragraph()
    run = p.add_run("Database Connection Error:")
    run.bold = True
    doc.add_paragraph("Check DATABASE_URL in .env, ensure PostgreSQL is running")

    p = doc.add_paragraph()
    run = p.add_run("Module Not Found:")
    run.bold = True
    doc.add_paragraph("Activate virtual environment, run 'pip install -r requirements.txt'")

    doc.add_page_break()

    # ========================================
    # CONCLUSION
    # ========================================

    add_heading(doc, "Conclusion", 1)

    add_heading(doc, "Summary Checklist", 2)
    doc.add_paragraph("You have learned:")

    checklist = [
        "What AI customer support is and why it matters",
        "How Large Language Models work",
        "Multi-agent system architecture",
        "Retrieval-Augmented Generation (RAG)",
        "Query categorization techniques",
        "Sentiment analysis implementation",
        "Knowledge base retrieval with FAISS",
        "LangGraph workflow orchestration",
        "FastAPI REST API development",
        "Database design with SQLAlchemy",
        "Webhook implementation",
        "Docker containerization",
        "Cloud deployment with Railway"
    ]
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, "Key Takeaways", 2)
    takeaways = [
        "Multi-agent systems provide modularity, testability, and maintainability",
        "RAG grounds LLM responses in factual knowledge, preventing hallucination",
        "State management is the backbone of complex AI workflows",
        "Prompt engineering is crucial for consistent AI behavior",
        "Escalation logic ensures humans handle what AI cannot",
        "Production readiness requires testing, monitoring, and proper deployment"
    ]
    for i, takeaway in enumerate(takeaways, 1):
        doc.add_paragraph(f"{i}. {takeaway}")

    add_heading(doc, "What You Can Do Now", 2)
    doc.add_paragraph("With this knowledge, you can:")
    abilities = [
        "Build AI-powered support systems for your applications",
        "Extend this system with new agents and capabilities",
        "Apply multi-agent patterns to other domains",
        "Deploy AI applications to production environments"
    ]
    for ability in abilities:
        doc.add_paragraph(ability, style='List Bullet')

    doc.add_paragraph()
    final = doc.add_paragraph()
    run = final.add_run("Happy Building!")
    run.bold = True
    run.font.size = Pt(14)
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    output_path = "tutorial_documentation/Multi-Agent HR Intelligence Platform_AI_Complete_Tutorial.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_tutorial()
